# -*- coding: utf-8 -*-

"""
Lab Work #7: Decision Support and Improving Consistency of Expert Evaluations
Variant #1

Program for analyzing pairwise comparison matrices, calculating alternative weights, and evaluating consistency of expert evaluations.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext, colorchooser, simpledialog
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from mpl_toolkits.mplot3d import Axes3D
import pandas as pd
import os
import io
import json
from decimal import Decimal
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import sys
from datetime import datetime
from decision_support import *

class DecisionSupportApp:
    """Class for the decision support application graphical interface."""

    def __init__(self, root):
        """Application initialization."""
        self.root = root
        self.root.title("Decision Support (Variant 1)")
        self.root.geometry("1400x900")

        # Style settings
        self.style = ttk.Style()
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure("TButton", padding=6)
        self.style.configure("TLabel", background="#f0f0f0", font=("Arial", 10))
        self.style.configure("Header.TLabel", font=("Arial", 12, "bold"))

        # Color themes
        self.themes = {
            "light": {"bg": "#f0f0f0", "fg": "black", "highlight": "blue"},
            "dark": {"bg": "#2b2b2b", "fg": "white", "highlight": "cyan"}
        }
        self.current_theme = "light"

        # State variables
        self.matrix_size = tk.IntVar(value=4)
        self.alternative_names = []
        self.matrix_entries = []
        self.current_matrix = None
        self.results = None
        self.sensitivity_results = None
        self.session_file = None

        # Added colors for visualization
        self.color_map = plt.cm.Set3
        self.highlight_color = 'red'

        # Operation log for undo
        self.operation_history = []
        self.max_history_size = 50

        # Create interface
        self.create_widgets()

        # Load example data for variant 1
        self.root.after(100, self.load_example_data)

    def create_widgets(self):
        """Create interface elements."""

        # Create main frames
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Left panel for matrix input
        left_frame = ttk.LabelFrame(main_frame, text="Pairwise Comparison Matrix")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Right panel for results
        right_frame = ttk.Frame(main_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Control panel
        control_frame = ttk.LabelFrame(right_frame, text="Control")
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        # Results panel
        results_frame = ttk.LabelFrame(right_frame, text="Analysis Results")
        results_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))

        # --- Left panel: pairwise comparison matrix ---
        # Matrix settings frame
        matrix_settings_frame = ttk.Frame(left_frame)
        matrix_settings_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(matrix_settings_frame, text="Matrix size:").pack(side=tk.LEFT, padx=(0, 5))
        size_spinbox = ttk.Spinbox(matrix_settings_frame, from_=2, to=10, textvariable=self.matrix_size, width=5, command=self.update_matrix_size)
        size_spinbox.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(matrix_settings_frame, text="Update matrix", command=self.update_matrix_size).pack(side=tk.LEFT)

        # Frame for alternative names
        names_frame = ttk.Frame(left_frame)
        names_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(names_frame, text="Alternative names:").pack(anchor=tk.W, padx=5, pady=(0, 5))

        self.names_entries_frame = ttk.Frame(names_frame)
        self.names_entries_frame.pack(fill=tk.X)

        # Frame for the matrix itself
        self.matrix_frame = ttk.Frame(left_frame)
        self.matrix_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create default matrix (4x4)
        self.create_matrix_entries(4)

        # --- Control panel ---
        # Control buttons
        btn_frame = ttk.Frame(control_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(btn_frame, text="Analyze", command=self.analyze_matrix, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Load example", command=self.load_example_data, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Export results", command=self.export_results, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="About", command=self.show_about, width=10).pack(side=tk.RIGHT, padx=2)

        # Additional function buttons
        extra_btn_frame = ttk.Frame(control_frame)
        extra_btn_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(extra_btn_frame, text="Sensitivity analysis", command=self.analyze_sensitivity, width=20).pack(side=tk.LEFT, padx=2)
        ttk.Button(extra_btn_frame, text="Adjust inconsistencies", command=self.adjust_inconsistencies, width=22).pack(side=tk.LEFT, padx=2)
        ttk.Button(extra_btn_frame, text="Generate consistent matrix", command=self.generate_consistent_matrix, width=28).pack(side=tk.LEFT, padx=2)

        # Import/export buttons
        import_export_frame = ttk.Frame(control_frame)
        import_export_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(import_export_frame, text="Save session", command=self.save_session).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Load session", command=self.load_session).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Import from CSV", command=self.import_from_csv).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Import from Excel", command=self.import_from_excel).pack(side=tk.LEFT, padx=2)

        # Theme toggle button
        ttk.Button(import_export_frame, text="Toggle theme", command=self.toggle_theme).pack(side=tk.RIGHT, padx=2)

        # --- Results panel ---
        # Tabs for results
        self.results_notebook = ttk.Notebook(results_frame)
        self.results_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Tab: alternative weights
        self.weights_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.weights_frame, text="Alternative Weights")
        self.create_weights_tab()

        # Tab: method comparison
        self.comparison_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.comparison_frame, text="Method Comparison")
        self.create_comparison_tab()

        # Tab: consistency
        self.consistency_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.consistency_frame, text="Consistency")
        self.create_consistency_tab()

        # Tab: inconsistent pairs
        self.inconsistent_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.inconsistent_frame, text="Inconsistent Pairs")
        self.create_inconsistent_pairs_tab()

        # Tab: 3D visualization
        self.visualization_3d_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.visualization_3d_frame, text="3D Visualization")
        self.create_3d_visualization_tab()

        # Tab: execution log
        self.log_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.log_frame, text="Execution Log")
        self.create_log_tab()

        # Add context help to all elements
        self.add_context_help()

    def add_context_help(self):
        """Add context help to all interface elements."""
        help_texts = {
            "matrix_size": "Size of the pairwise comparison matrix (from 2 to 10)",
            "alternative_names": "Names of the alternatives being compared",
            "analyze_button": "Perform matrix analysis and calculate alternative weights",
            "sensitivity_button": "Analyze the sensitivity of results to changes in the matrix",
            "adjust_button": "Automatically correct the most inconsistent pairwise comparisons",
            "save_session": "Save the entire current analysis session (matrix, settings, results)",
            "load_session": "Load a previously saved analysis session",
            "weights_tab": "View alternative weights calculated by various methods",
            "comparison_tab": "Comparison of results obtained by different methods",
            "consistency_tab": "Consistency indicators of expert evaluations",
            "inconsistent_tab": "List of the most inconsistent pairwise comparisons",
            "3d_tab": "3D visualization of the pairwise comparison matrix and alternative weights",
            "log_tab": "Log of performed operations and system messages"
        }

    def show_tooltip(self, text):
        """Display a tooltip."""
        try:
            self.tooltip = tk.Toplevel(self.root)
            self.tooltip.wm_overrideredirect(True)
            self.tooltip.wm_geometry(f"+{self.root.winfo_pointerx() + 10}+{self.root.winfo_pointery() + 10}")
            label = ttk.Label(self.tooltip, text=text, background="lightyellow", relief="solid", borderwidth=1)
            label.pack()
        except Exception:
            pass

    def hide_tooltip(self):
        """Hide the tooltip."""
        try:
            if hasattr(self, 'tooltip'):
                self.tooltip.destroy()
        except Exception:
            pass

    def create_matrix_entries(self, size):
        """Create input fields for a matrix of a given size."""
        # Clear the previous matrix
        for widget in self.matrix_frame.winfo_children():
            widget.destroy()
        self.matrix_entries = []
        self.alternative_names = []

        # Create column headers
        for j in range(size):
            label = ttk.Label(self.matrix_frame, text=f"Alternative {j + 1}", font=("Arial", 9, "bold"))
            label.grid(row=0, column=j + 1, padx=2, pady=2)

        # Create input fields for the matrix
        for i in range(size):
            # Row header
            label = ttk.Label(self.matrix_frame, text=f"Alternative {i + 1}", font=("Arial", 9, "bold"))
            label.grid(row=i + 1, column=0, padx=2, pady=2)

            row_entries = []

            for j in range(size):
                entry_var = tk.StringVar(value="1.0" if i == j else "3.0" if i < j else "0.33")
                entry = ttk.Entry(self.matrix_frame, textvariable=entry_var, width=8)

                # Diagonal elements are not editable (must be 1)
                if i == j:
                    entry.config(state="readonly")

                # Add input validation
                entry.bind("<FocusOut>", lambda event, r=i, c=j: self.validate_entry(r, c))
                entry.bind("<Return>", lambda event, r=i, c=j: self.on_enter_pressed(r, c))
                entry.grid(row=i + 1, column=j + 1, padx=2, pady=2)

                row_entries.append(entry_var)

            self.matrix_entries.append(row_entries)

        # Update alternative names
        self.update_alternative_names(size)

    def validate_entry(self, i, j):
        """Check the correctness of the entered value."""
        if i == j:
            return

        try:
            value = float(self.matrix_entries[i][j].get())
            if value <= 0:
                raise ValueError("Value must be positive")

            # Check transitivity
            message = f"Element [{i + 1},{j + 1}] set to {value:.2f}"
            self.log_message(message)
        except ValueError as e:
            messagebox.showerror("Input error", f"Invalid value: {str(e)}")
            self.log_message(f"Input error in element [{i + 1},{j + 1}]")
            self.matrix_entries[i][j].set("1.0")

    def on_enter_pressed(self, i, j):
        """Action when Enter is pressed."""
        self.log_message(f"Value changed in [{i + 1},{j + 1}]")
        # Add to operation history
        self.add_to_operation_history(f"Change element [{i + 1},{j + 1}]")

    def update_alternative_names(self, size):
        """Update fields for alternative names."""
        # Clear previous fields
        for widget in self.names_entries_frame.winfo_children():
            widget.destroy()
        self.alternative_names = []

        for i in range(size):
            frame = ttk.Frame(self.names_entries_frame)
            frame.pack(fill=tk.X, pady=2)

            ttk.Label(frame, text=f"Alternative {i + 1}:").pack(side=tk.LEFT, padx=(0, 5))
            name_var = tk.StringVar(value=f"Option {i + 1}")
            entry = ttk.Entry(frame, textvariable=name_var, width=20)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
            self.alternative_names.append(name_var)

    def update_matrix_size(self):
        """Update the matrix size."""
        size = self.matrix_size.get()
        self.create_matrix_entries(size)
        self.update_alternative_names(size)

    def get_matrix_from_entries(self):
        """Get the matrix from the input fields."""
        size = len(self.matrix_entries)
        matrix = np.ones((size, size))

        try:
            for i in range(size):
                for j in range(size):
                    if i != j:
                        val = float(self.matrix_entries[i][j].get())
                        # Check for positivity
                        if val <= 0:
                            raise ValueError("Matrix values must be positive")
                        matrix[i, j] = val
                        # For symmetric elements (j, i) should be 1/val
                        matrix[j, i] = 1 / val
            return matrix
        except ValueError as e:
            messagebox.showerror("Input error", f"Invalid value in the matrix: {str(e)}")
            return None

    def load_example_data(self):
        """Load example data for variant 1."""
        # For variant 1, use an example with choosing a place for building an object
        size = 4
        self.matrix_size.set(size)
        self.create_matrix_entries(size)
        self.update_alternative_names(size)
        # Set alternative names according to variant #1
        alternative_names = ["Project Management", "Requirements Analysis", "Design", "Implementation"]
        for i, name in enumerate(alternative_names):
            if i < len(self.alternative_names):
                self.alternative_names[i].set(name)
        # Fill the matrix with an example
        example_matrix = [
            [1, 3, 5, 7],
            [1 / 3, 1, 3, 5],
            [1 / 5, 1 / 3, 1, 3],
            [1 / 7, 1 / 5, 1 / 3, 1]
        ]
        for i in range(size):
            for j in range(size):
                if i != j:
                    self.matrix_entries[i][j].set(f"{example_matrix[i][j]:.2f}")

        self.log_message("Example data for variant 1 loaded: information systems development.")
        self.log_message("Pairwise comparison matrix filled with sample values.")

        # Perform analysis after loading the example
        self.analyze_matrix()

    def analyze_matrix(self):
        """Analyze the pairwise comparison matrix."""
        matrix = self.get_matrix_from_entries()
        if matrix is None:
            return

        self.current_matrix = matrix

        try:
            # Perform consistency analysis
            self.results = analyze_consistency(matrix)

            # Update all tabs
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.update_3d_visualization()

            self.log_message("Matrix analysis completed successfully.")
        except Exception as e:
            messagebox.showerror("Analysis error", f"An error occurred during matrix analysis: {str(e)}")
            self.log_message(f"Error: {str(e)}")

    def analyze_sensitivity(self):
        """Sensitivity analysis of the matrix."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Sensitivity analysis", "First, perform matrix analysis.")
            return

        try:
            # Perform sensitivity analysis
            self.sensitivity_results = create_sensitivity_analysis(
                self.current_matrix, self.results, perturbation_factor=0.1
            )

            # Display results
            i, j = self.sensitivity_results['most_sensitive_pair']
            names = [var.get() for var in self.alternative_names]

            message = (
                f"Most sensitive pair: {names[i]} vs {names[j]} \n"
                f"Maximum sensitivity: {self.sensitivity_results['max_sensitivity']:.6f} \n"
                f"Perturbation factor: {self.sensitivity_results['perturbation_factor'] * 100:.1f} %"
            )
            messagebox.showinfo("Sensitivity analysis", message)
            self.log_message(f"Sensitivity analysis performed: {message}")

            # Update 3D visualization
            self.update_3d_visualization()
        except Exception as e:
            messagebox.showerror("Analysis error", f"An error occurred during sensitivity analysis: {str(e)}")
            self.log_message(f"Sensitivity analysis error: {str(e)}")

    def adjust_inconsistencies(self):
        """Adjust inconsistent pairwise comparisons."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Adjust inconsistencies", "First, perform matrix analysis.")
            return

        try:
            # Perform adjustment
            corrected_matrix, changes = adjust_inconsistent_pairs(
                self.current_matrix, self.results['inconsistent_pairs'], threshold=0.1
            )

            if not changes:
                messagebox.showinfo("Adjustment", "Inconsistency is within acceptable limits.")
                self.log_message("Adjusting inconsistencies: no changes required.")
                return

            # Ask the user to apply changes
            confirm = messagebox.askyesno(
                "Adjustment",
                f"{len(changes)} inconsistent pairs found.\nApply automatic corrections?"
            )

            if confirm:
                # Apply changes to the interface
                for change in changes:
                    i, j = change['pair']
                    self.matrix_entries[i][j].set(f"{change['corrected']:.2f}")
                    self.matrix_entries[j][i].set(f"{1 / change['corrected']:.2f}")
                # Update the matrix and restart analysis
                self.current_matrix = corrected_matrix
                self.results = analyze_consistency(corrected_matrix)

                # Update all tabs
                self.update_weights_tab()
                self.update_comparison_tab()
                self.update_consistency_tab()
                self.update_inconsistent_pairs_tab()
                self.update_3d_visualization()

                self.log_message(f"Automatic corrections applied for {len(changes)} inconsistent pairs.")
                self.log_message("The matrix has been corrected and reanalyzed.")
        except Exception as e:
            messagebox.showerror("Adjustment error", f"An error occurred during adjustment: {str(e)}")
            self.log_message(f"Adjustment error: {str(e)}")

    def generate_consistent_matrix(self):
        """Generate a consistent matrix."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Generate matrix", "First, perform matrix analysis.")
            return

        try:
            # Get the consistent matrix from the results
            consistent_matrix = self.results['consistent_matrix']

            # Create a new window to display the consistent matrix
            dialog = tk.Toplevel(self.root)
            dialog.title("Consistent Matrix")
            dialog.geometry("600x500")

            # Create a frame for the matrix
            matrix_frame = ttk.Frame(dialog)
            matrix_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            # Create a table
            columns = [f"Alternative {i + 1}" for i in range(consistent_matrix.shape[0])]
            tree = ttk.Treeview(matrix_frame, columns=columns, show="headings", height=10)

            # Configure headers
            names = [var.get() for var in self.alternative_names]
            for i, name in enumerate(names):
                tree.heading(columns[i], text=name)
                tree.column(columns[i], width=100)

            # Add data
            for i, row in enumerate(consistent_matrix):
                values = [f"{val:.3f}" for val in row]
                tree.insert("", tk.END, values=values)

            # Add scrollbars
            v_scrollbar = ttk.Scrollbar(matrix_frame, orient=tk.VERTICAL, command=tree.yview)
            h_scrollbar = ttk.Scrollbar(matrix_frame, orient=tk.HORIZONTAL, command=tree.xview)
            tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

            tree.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

            # Control buttons
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(fill=tk.X, padx=10, pady=10)

            ttk.Button(btn_frame, text="Copy", command=lambda: self.copy_matrix_to_clipboard(consistent_matrix)).pack(side=tk.LEFT)
            ttk.Button(btn_frame, text="Apply", command=lambda: self.apply_consistent_matrix(consistent_matrix, dialog)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Close", command=dialog.destroy).pack(side=tk.RIGHT)

            self.log_message("Consistent matrix generated based on current weights.")
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred while generating the consistent matrix: {str(e)}")
            self.log_message(f"Error generating consistent matrix: {str(e)}")

    def copy_matrix_to_clipboard(self, matrix):
        """Copy the matrix to the clipboard."""
        try:
            # Convert the matrix to text format
            text = ""
            names = [var.get() for var in self.alternative_names]

            # Add headers
            text += "\t" + "\t".join(names) + "\n"

            # Add matrix rows
            for i, row in enumerate(matrix):
                text += names[i] + "\t" + "\t".join([f"{val:.3f}" for val in row]) + "\n"

            self.root.clipboard_clear()
            self.root.clipboard_append(text)

            self.log_message("Consistent matrix copied to clipboard.")
            messagebox.showinfo("Copy", "Matrix copied to clipboard.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to copy matrix: {str(e)}")
            self.log_message(f"Error copying matrix: {str(e)}")

    def apply_consistent_matrix(self, matrix, dialog):
        """Apply the consistent matrix to the main interface."""
        try:
            # Apply values to the main interface
            for i in range(matrix.shape[0]):
                for j in range(matrix.shape[1]):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{matrix[i, j]:.3f}")

            # Update the main matrix
            self.current_matrix = matrix
            self.results = analyze_consistency(matrix)

            # Update all tabs
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.update_3d_visualization()

            # Close the dialog window
            dialog.destroy()

            self.log_message("Consistent matrix applied to the main interface.")
            messagebox.showinfo("Apply", "Consistent matrix successfully applied.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to apply matrix: {str(e)}")
            self.log_message(f"Error applying matrix: {str(e)}")

    def create_weights_tab(self):
        """Create a tab for displaying alternative weights."""
        # Frame for method selection
        method_frame = ttk.Frame(self.weights_frame)
        method_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(method_frame, text="Weight calculation method:").pack(side=tk.LEFT, padx=(0, 10))
        self.weights_method = tk.StringVar(value="eigenvector")
        methods = [
            ("Eigenvector method", "eigenvector"),
            ("Logarithmic least squares method", "log_least_squares"),
            ("Geometric mean method", "geometric_mean"),
            ("Line method", "line_method"),
            ("AHP method", "ahp")
        ]

        for text, method in methods:
            rb = ttk.Radiobutton(method_frame, text=text, value=method, variable=self.weights_method, command=self.update_weights_display)
            rb.pack(side=tk.LEFT, padx=5)

        # Frame for displaying weights
        self.weights_display_frame = ttk.Frame(self.weights_frame)
        self.weights_display_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create chart and table
        self.create_weights_display()

    def create_weights_display(self):
        """Create elements for displaying weights."""
        # Clear previous content
        for widget in self.weights_display_frame.winfo_children():
            widget.destroy()

        # Create frame for chart
        chart_frame = ttk.Frame(self.weights_display_frame)
        chart_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Create chart
        self.weights_fig, self.weights_ax = plt.subplots(figsize=(6, 4))
        self.weights_canvas = FigureCanvasTkAgg(self.weights_fig, master=chart_frame)
        self.weights_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Create frame for table
        table_frame = ttk.Frame(self.weights_display_frame)
        table_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Create table
        columns = ("alternative", "weight")
        self.weights_tree = ttk.Treeview(table_frame, columns=columns, show="headings")

        # Configure headers
        self.weights_tree.heading("alternative", text="Alternative")
        self.weights_tree.heading("weight", text="Weight")

        # Configure columns
        self.weights_tree.column("alternative", width=150)
        self.weights_tree.column("weight", width=100, anchor=tk.E)

        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.weights_tree.yview)
        self.weights_tree.configure(yscrollcommand=scrollbar.set)
        self.weights_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def update_weights_tab(self):
        """Update the tab with alternative weights."""
        if self.results is None:
            return

        # Update weight display
        self.update_weights_display()

    def update_weights_display(self):
        """Update weight display based on the selected method."""
        if self.results is None:
            return

        method = self.weights_method.get()

        # Determine which weights to use
        if method == "eigenvector":
            weights = self.results['weights_eigenvector']
            method_name = "Eigenvector method"
        elif method == "log_least_squares":
            weights = self.results['weights_log_least_squares']
            method_name = "Logarithmic least squares method"
        elif method == "geometric_mean":
            weights = self.results['weights_geometric_mean']
            method_name = "Geometric mean method"
        elif method == "ahp":
            weights, _, _, _ = calculate_weights_ahp(self.current_matrix)
            method_name = "AHP method"
        else: # line_method
            weights = self.results['weights_line_method']
            method_name = "Line method"

        # Update chart
        self.weights_ax.clear()

        # Get alternative names
        names = [var.get() for var in self.alternative_names]

        # Plot bar chart
        bars = self.weights_ax.bar(names, weights, color='skyblue')
        self.weights_ax.set_title(f'Alternative weights ({method_name})')
        self.weights_ax.set_ylabel('Weight')
        self.weights_ax.set_ylim(0, max(weights) * 1.2)

        # Add values above bars
        for bar in bars:
            height = bar.get_height()
            self.weights_ax.annotate(f'{height:.4f}', xy=(bar.get_x() + bar.get_width() / 2, height), xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')

        self.weights_ax.tick_params(axis='x', rotation=15)
        self.weights_fig.tight_layout()
        self.weights_canvas.draw()

        # Update table
        for item in self.weights_tree.get_children():
            self.weights_tree.delete(item)

        for i, (name, weight) in enumerate(zip(names, weights)):
            self.weights_tree.insert("", tk.END, values=(name, f"{weight:.6f}"))

        # Determine the best alternative
        best_idx = np.argmax(weights)
        best_name = names[best_idx]
        best_weight = weights[best_idx]

        self.log_message(f"Best alternative by {method_name}: {best_name} (weight = {best_weight:.4f})")

    def create_comparison_tab(self):
        """Create a tab for method comparison."""
        # Frame for displaying comparison
        self.comparison_display_frame = ttk.Frame(self.comparison_frame)
        self.comparison_display_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create chart and table
        self.create_comparison_display()

    def create_comparison_display(self):
        """Create elements for displaying method comparison."""
        # Clear previous content
        for widget in self.comparison_display_frame.winfo_children():
            widget.destroy()

        # Create frame for chart
        chart_frame = ttk.Frame(self.comparison_display_frame)
        chart_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Create comparison chart
        self.comparison_fig, self.comparison_ax = plt.subplots(figsize=(6, 4))
        self.comparison_canvas = FigureCanvasTkAgg(self.comparison_fig, master=chart_frame)
        self.comparison_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Create frame for table
        table_frame = ttk.Frame(self.comparison_display_frame)
        table_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Create comparison table
        columns = ("alternative", "eigenvector", "log", "geometric", "line", "ahp")
        self.comparison_tree = ttk.Treeview(table_frame, columns=columns, show="headings")

        # Configure headers
        self.comparison_tree.heading("alternative", text="Alternative")
        self.comparison_tree.heading("eigenvector", text="Eigenvector")
        self.comparison_tree.heading("log", text="Log.\nleast sq.")
        self.comparison_tree.heading("geometric", text="Geometric\nmean")
        self.comparison_tree.heading("line", text="Line\nmethod")
        self.comparison_tree.heading("ahp", text="AHP\nmethod")

        # Configure columns
        self.comparison_tree.column("alternative", width=120)
        self.comparison_tree.column("eigenvector", width=80, anchor=tk.E)
        self.comparison_tree.column("log", width=80, anchor=tk.E)
        self.comparison_tree.column("geometric", width=80, anchor=tk.E)
        self.comparison_tree.column("line", width=80, anchor=tk.E)
        self.comparison_tree.column("ahp", width=80, anchor=tk.E)

        # Add scrollbar
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.comparison_tree.yview)
        self.comparison_tree.configure(yscrollcommand=scrollbar.set)
        self.comparison_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def update_comparison_tab(self):
        """Update the method comparison tab."""
        if self.results is None:
            return

        # Update chart
        self.comparison_ax.clear()

        # Get alternative names
        names = [var.get() for var in self.alternative_names]

        # Prepare data for chart
        weights_eig = self.results['weights_eigenvector']
        weights_log = self.results['weights_log_least_squares']
        weights_geo = self.results['weights_geometric_mean']
        weights_line = self.results['weights_line_method']
        weights_ahp, _, _, _ = calculate_weights_ahp(self.current_matrix)

        # Bar width
        bar_width = 0.15
        indices = np.arange(len(names))

        # Plot bar chart
        self.comparison_ax.bar(indices - 2 * bar_width, weights_eig, bar_width, label='Eigenvector', alpha=0.8)
        self.comparison_ax.bar(indices - bar_width, weights_log, bar_width, label='Log. least sq.', alpha=0.8)
        self.comparison_ax.bar(indices, weights_geo, bar_width, label='Geometric mean', alpha=0.8)
        self.comparison_ax.bar(indices + bar_width, weights_line, bar_width, label='Line method', alpha=0.8)
        self.comparison_ax.bar(indices + 2 * bar_width, weights_ahp, bar_width, label='AHP', alpha=0.8)

        self.comparison_ax.set_xlabel('Alternatives')
        self.comparison_ax.set_ylabel('Weights')
        self.comparison_ax.set_title('Comparison of weight calculation methods')
        self.comparison_ax.set_xticks(indices)
        self.comparison_ax.set_xticklabels(names, rotation=15)
        self.comparison_ax.legend()

        self.comparison_fig.tight_layout()
        self.comparison_canvas.draw()

        # Update table
        for item in self.comparison_tree.get_children():
            self.comparison_tree.delete(item)

        for i, name in enumerate(names):
            self.comparison_tree.insert("", tk.END, values=(
                name,
                f"{weights_eig[i]:.4f}", f"{weights_log[i]:.4f}", f"{weights_geo[i]:.4f}", f"{weights_line[i]:.4f}", f"{weights_ahp[i]:.4f}"
            ))

        self.log_message("Method comparison table updated.")

    def create_consistency_tab(self):
        """Create a tab for displaying consistency indicators."""
        # Frame for main indicators
        main_frame = ttk.Frame(self.consistency_frame)
        main_frame.pack(fill=tk.X, padx=5, pady=5)

        # Main indicators table
        self.consistency_tree = ttk.Treeview(main_frame, columns=("metric", "value"), show="headings", height=8)
        self.consistency_tree.heading("metric", text="Metric")
        self.consistency_tree.heading("value", text="Value")
        self.consistency_tree.column("metric", width=250)
        self.consistency_tree.column("value", width=150, anchor=tk.E)
        self.consistency_tree.pack(fill=tk.X, padx=5, pady=5)

        # Frame for correlation chart
        chart_frame = ttk.Frame(self.consistency_frame)
        chart_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create correlation chart
        self.corr_fig, self.corr_ax = plt.subplots(figsize=(6, 4))
        self.corr_canvas = FigureCanvasTkAgg(self.corr_fig, master=chart_frame)
        self.corr_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def update_consistency_tab(self):
        """Update the consistency tab."""
        if self.results is None:
            return

        # Update metrics table
        for item in self.consistency_tree.get_children():
            self.consistency_tree.delete(item)

        metrics = [
            ("Eigenvalue λ_max", f"{self.results['lambda_max']:.4f}"),
            ("Consistency index CI", f"{self.results['CI']:.4f}"),
            ("Consistency ratio CR", f"{self.results['CR']:.4f}"),
            ("Average correlation", f"{self.results['mean_correlation']:.4f}"),
            ("Chi-square statistic", f"{self.results['chi2_statistic']:.4f}"),
            ("p-value for Chi-square", f"{self.results['chi2_p_value']:.4f}"),
            ("Degrees of freedom", f"{self.results['chi2_df']}"),
            ("Matrix size", f"{self.current_matrix.shape[0]} x {self.current_matrix.shape[1]}")
        ]

        for metric, value in metrics:
            self.consistency_tree.insert("", tk.END, values=(metric, value))

        # Consistency assessment
        cr_threshold = 0.1
        chi2_threshold = 0.05

        if self.results['CR'] < cr_threshold and self.results['chi2_p_value'] > chi2_threshold:
            self.log_message("The pairwise comparison matrix is considered consistent (CR < 0.1 and p-value > 0.05).")
        else:
            self.log_message("The pairwise comparison matrix has inconsistency (check CR and p-value).")

        # Update correlation chart
        self.corr_ax.clear()

        # Get correlation matrix
        corr_matrix = self.results['correlation_matrix']
        size = corr_matrix.shape[0]

        # Plot heatmap
        im = self.corr_ax.imshow(corr_matrix, cmap='coolwarm', vmin=-1, vmax=1)

        # Add alternative names
        names = [var.get() for var in self.alternative_names]
        self.corr_ax.set_xticks(range(size))
        self.corr_ax.set_yticks(range(size))
        self.corr_ax.set_xticklabels(names, rotation=45, ha='right')
        self.corr_ax.set_yticklabels(names)

        # Add values to cells
        for i in range(size):
            for j in range(size):
                self.corr_ax.text(j, i, f"{corr_matrix[i, j]:.2f}", ha="center", va="center", color="w" if abs(corr_matrix[i, j]) > 0.5 else "black")

        self.corr_ax.set_title('Correlation matrix between alternatives')
        self.corr_fig.colorbar(im, ax=self.corr_ax)
        self.corr_fig.tight_layout()
        self.corr_canvas.draw()

    def create_inconsistent_pairs_tab(self):
        """Create a tab for displaying inconsistent pairs."""
        # Filter frame
        filter_frame = ttk.Frame(self.inconsistent_frame)
        filter_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(filter_frame, text="Filter by inconsistency degree >").pack(side=tk.LEFT)
        self.inconsistency_threshold = tk.DoubleVar(value=0.1)
        threshold_entry = ttk.Entry(filter_frame, textvariable=self.inconsistency_threshold, width=8)
        threshold_entry.pack(side=tk.LEFT, padx=(5, 10))

        ttk.Button(filter_frame, text="Filter", command=self.update_inconsistent_pairs_tab).pack(side=tk.LEFT)
        ttk.Button(filter_frame, text="Reset", command=lambda: self.inconsistency_threshold.set(0.1)).pack(side=tk.LEFT, padx=(5, 0))

        # Inconsistent pairs table
        self.inconsistent_tree = ttk.Treeview(self.inconsistent_frame, columns=("triplet", "inconsistency", "recommendation"), show="headings", height=10)
        self.inconsistent_tree.heading("triplet", text="Alternative triplet")
        self.inconsistent_tree.heading("inconsistency", text="Inconsistency degree")
        self.inconsistent_tree.heading("recommendation", text="Recommended value")
        self.inconsistent_tree.column("triplet", width=200)
        self.inconsistent_tree.column("inconsistency", width=150, anchor=tk.E)
        self.inconsistent_tree.column("recommendation", width=150, anchor=tk.E)
        self.inconsistent_tree.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Recommendations frame
        rec_frame = ttk.LabelFrame(self.inconsistent_frame, text="Recommendations")
        rec_frame.pack(fill=tk.X, padx=5, pady=5)

        self.recommendations_text = scrolledtext.ScrolledText(rec_frame, height=5, wrap=tk.WORD)
        self.recommendations_text.pack(fill=tk.X, padx=5, pady=5)
        self.recommendations_text.config(state=tk.DISABLED)

    def update_inconsistent_pairs_tab(self):
        """Update the tab with inconsistent pairs."""
        if self.results is None:
            return

        # Update table
        for item in self.inconsistent_tree.get_children():
            self.inconsistent_tree.delete(item)

        names = [var.get() for var in self.alternative_names]
        threshold = self.inconsistency_threshold.get()

        filtered_pairs = [(a, b, c, inc) for a, b, c, inc in self.results['inconsistent_pairs'] if inc > threshold]

        for i, (a, b, c, inc) in enumerate(filtered_pairs):
            triplet = f"{names[a]} → {names[b]} → {names[c]}"
            # Calculate recommended value
            recommended = self.current_matrix[a, b] * self.current_matrix[b, c]

            self.inconsistent_tree.insert("", tk.END, values=(triplet, f"{inc:.4f}", f"{recommended:.2f}"))

        # Update recommendations
        self.recommendations_text.config(state=tk.NORMAL)
        self.recommendations_text.delete(1.0, tk.END)

        if filtered_pairs:
            a, b, c, inc = filtered_pairs[0]

            names = [var.get() for var in self.alternative_names]

            rec_text = (
                f"Most inconsistent triplet: {names[a]} → {names[b]} → {names[c]} \n\n"
                "Recommendations:\n"
                f"1. Check the comparison evaluation of {names[a]} and {names[b]} (value: {self.current_matrix[a, b]:.2f})\n"
                f"2. Check the comparison evaluation of {names[b]} and {names[c]} (value: {self.current_matrix[b, c]:.2f})\n"
                f"3. Check the comparison evaluation of {names[a]} and {names[c]} (value: {self.current_matrix[a, c]:.2f})\n"
                f"4. For consistency, the value of {names[a]} and {names[c]} should be close to {self.current_matrix[a, b] * self.current_matrix[b, c]:.2f}\n\n"
                "Possible actions:\n"
                "- Review expert evaluations for the specified pairs\n"
                "- Clarify comparison criteria\n"
                "- Conduct additional discussion with experts"
            )

            self.recommendations_text.insert(tk.END, rec_text)
        else:
            self.recommendations_text.insert(tk.END, "No inconsistent pairs above the threshold found.")

        self.recommendations_text.config(state=tk.DISABLED)

    def create_3d_visualization_tab(self):
        """Create a tab for 3D visualization."""
        # Frame for 3D chart
        self.vis_3d_frame = ttk.Frame(self.visualization_3d_frame)
        self.vis_3d_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Create 3D chart
        self.fig_3d = plt.figure(figsize=(8,6))
        self.ax_3d = self.fig_3d.add_subplot(111, projection='3d')
        self.canvas_3d = FigureCanvasTkAgg(self.fig_3d, master=self.vis_3d_frame)
        self.canvas_3d.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Control frame
        control_frame = ttk.Frame(self.visualization_3d_frame)
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(control_frame, text="Update 3D chart", command=self.update_3d_visualization).pack(side=tk.LEFT)
        ttk.Button(control_frame, text="Change view", command=self.rotate_view).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="Export 3D", command=self.export_3d_plot).pack(side=tk.LEFT, padx=5)

        # Initialize view angle
        self.elev = 30
        self.azim = 45

    def rotate_view(self):
        """Change the viewing angle of the 3D chart."""
        self.elev = (self.elev + 30) % 360
        self.azim = (self.azim + 45) % 360
        self.update_3d_visualization()

    def export_3d_plot(self):
        """Export the 3D chart to a file."""
        if self.results is None:
            messagebox.showwarning("Export", "First, perform matrix analysis.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("PDF files", "*.pdf"), ("SVG files", "*.svg")],
            title="Export 3D chart"
        )

        if file_path:
            try:
                self.fig_3d.savefig(file_path, dpi=300, bbox_inches='tight')
                self.log_message(f"3D chart exported to file: {os.path.basename(file_path)}")
                messagebox.showinfo("Export", f"3D chart successfully saved to file:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Export error", f"Failed to export 3D chart: {str(e)}")
                self.log_message(f"Error exporting 3D chart: {str(e)}")

    def update_3d_visualization(self):
        """Update 3D visualization."""
        if self.results is None:
            return

        # Clear previous chart
        self.ax_3d.clear()

        # Get alternative names
        names = [var.get() for var in self.alternative_names]
        n = len(names)

        # Create data for 3D chart
        x_pos = np.arange(n)
        y_pos = np.arange(n)
        x_pos, y_pos = np.meshgrid(x_pos, y_pos)

        # Use the pairwise comparison matrix as z-values
        z_values = self.current_matrix

        # Plot surface
        surf = self.ax_3d.plot_surface(x_pos, y_pos, z_values, cmap='viridis', alpha=0.8)

        # Add bars for weights
        weights = self.results['weights_eigenvector']

        for i, w in enumerate(weights):
            self.ax_3d.bar3d(i, n + 1, 0, 0.5, 0.5, w, color=self.color_map(i % 12), alpha=0.8)

        # Configure axes
        self.ax_3d.set_xlabel('Alternatives (columns)')
        self.ax_3d.set_ylabel('Alternatives (rows)')
        self.ax_3d.set_zlabel('Value')

        # Set labels
        self.ax_3d.set_xticks(range(n))
        self.ax_3d.set_xticklabels(names, rotation=45)
        self.ax_3d.set_yticks(range(n))
        self.ax_3d.set_yticklabels(names, rotation=-45)

        # Add legend for weights
        self.ax_3d.text2D(0.05, 0.95, "Alternative weights:", transform=self.ax_3d.transAxes, fontsize=10)
        for i, (name, w) in enumerate(zip(names, weights)):
            self.ax_3d.text2D(0.05, 0.90 - i * 0.05, f"{name}: {w:.3f}", transform=self.ax_3d.transAxes, fontsize=9)

        # Set view angle
        self.ax_3d.view_init(elev=self.elev, azim=self.azim)

        # Add color scale
        self.fig_3d.colorbar(surf, ax=self.ax_3d, shrink=0.5, aspect=5)
        self.ax_3d.set_title('3D Visualization of Pairwise Comparison Matrix and Alternative Weights')
        self.fig_3d.tight_layout()
        self.canvas_3d.draw()

    def create_log_tab(self):
        """Create a tab for the execution log."""
        self.log_text = scrolledtext.ScrolledText(self.log_frame, wrap=tk.WORD, state=tk.DISABLED)
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Log control buttons
        btn_frame = ttk.Frame(self.log_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(btn_frame, text="Clear log", command=self.clear_log).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Copy log", command=self.copy_log).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(btn_frame, text="Save log to file", command=self.save_log).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(btn_frame, text="Export full report", command=self.export_full_report).pack(side=tk.RIGHT, padx=(5, 0))

        # Add undo/redo buttons
        undo_redo_frame = ttk.Frame(self.log_frame)
        undo_redo_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(undo_redo_frame, text="Undo (Ctrl+Z)", command=self.undo_operation).pack(side=tk.LEFT)
        ttk.Button(undo_redo_frame, text="Redo (Ctrl+Y)", command=self.redo_operation).pack(side=tk.LEFT, padx=(5, 0))

        # Bind keyboard shortcuts
        self.root.bind('<Control-z>', lambda e: self.undo_operation())
        self.root.bind('<Control-y>', lambda e: self.redo_operation())

    def log_message(self, message):
        """Add a message to the log."""
        try:
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)
        except Exception as e:
            print(f"[LOG ERROR] {message} | Exception: {e}")

    def clear_log(self):
        """Clear the log."""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)
        self.log_message("Log cleared.")

    def copy_log(self):
        """Copy the log to the clipboard."""
        log_content = self.log_text.get(1.0, tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(log_content)
        self.log_message("Log copied to clipboard.")

    def save_log(self):
        """Save the log to a file."""
        log_content = self.log_text.get(1.0, tk.END)
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="Save log"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(log_content)
                self.log_message(f"Log saved to file: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Save error", f"Failed to save log: {str(e)}")
                self.log_message(f"Error saving log: {str(e)}")

    def add_to_operation_history(self, operation):
        """Add an operation to the history."""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.operation_history.append({'timestamp': timestamp, 'operation': operation})

        # Limit history size
        if len(self.operation_history) > self.max_history_size:
            self.operation_history.pop(0)

    def undo_operation(self):
        """Undo the last operation."""
        if not self.operation_history:
            messagebox.showinfo("Undo", "No operations to undo.")
            return

        last_operation = self.operation_history.pop()
        self.log_message(f"Undo operation: {last_operation['operation']} (time: {last_operation['timestamp']})")

    def redo_operation(self):
        """Redo the undone operation."""
        messagebox.showinfo("Redo", "Redo function is not fully implemented yet.")

    def export_full_report(self):
        """Export a full report on the lab work."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Export report", "First, perform matrix analysis.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[
                ("Word documents", "*.docx"),
                ("PDF files", "*.pdf"),
                ("Text files", "*.txt")
            ],
            title="Export full report"
        )

        if not file_path:
            return

        try:
            # Create report
            report = self.generate_report()
            # Save depending on extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.docx':
                # Export to Word document
                self.export_to_docx(report, file_path)
            elif ext == '.pdf':
                # Export to PDF
                self.export_to_pdf(report, file_path)
            else: # .txt
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report)

            self.log_message(f"Full report exported to file: {os.path.basename(file_path)}")
            messagebox.showinfo("Export report", f"Full report successfully saved to file:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Export error", f"Failed to export report: {str(e)}")
            self.log_message(f"Error exporting report: {str(e)}")

    def generate_report(self):
        """Generate a full report on the lab work."""
        report = []
        report.append("=" * 60)
        report.append("REPORT ON LABORATORY WORK #7")
        report.append("Decision Support and Improving Consistency of Expert Evaluations")
        report.append("Variant #1: Information Systems Development")
        report.append("=" * 60)
        report.append(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append("\n")

        # Original matrix
        report.append("1. ORIGINAL PAIRWISE COMPARISON MATRIX")
        report.append("-" * 40)

        names = [var.get() for var in self.alternative_names]
        report.append(" " + " ".join([f"{name:>10}" for name in names]))

        for i, name in enumerate(names):
            row = [f"{name:>10}"]
            for j in range(len(names)):
                if i == j:
                    row.append(" 1.00")
                else:
                    row.append(f"{self.current_matrix[i][j]:>10.2f}")
            report.append("".join(row))

        report.append("\n")

        # Alternative weights
        report.append("2. ALTERNATIVE WEIGHTS")
        report.append("-" * 40)

        report.append(
            f"{'Alternative':<15} {'Eigenvector':<15} {'Log. least sq.':<15} {'Geometric mean':<15} {'Line method':<15} {'AHP method':<15}"
        )

        weights_eig = self.results['weights_eigenvector']
        weights_log = self.results['weights_log_least_squares']
        weights_geo = self.results['weights_geometric_mean']
        weights_line = self.results['weights_line_method']
        weights_ahp, _, _, _ = calculate_weights_ahp(self.current_matrix)

        for i, name in enumerate(names):
            report.append(
                f"{name:<15} {weights_eig[i]:<15.6f} {weights_log[i]:<15.6f} {weights_geo[i]:<15.6f} {weights_line[i]:<15.6f} {weights_ahp[i]:<15.6f}"
            )

        report.append("\n")

        # Consistency indicators
        report.append("3. CONSISTENCY INDICATORS")
        report.append("-" * 40)

        report.append(f"Eigenvalue λ_max: {self.results['lambda_max']:.4f}")
        report.append(f"Consistency index CI: {self.results['CI']:.4f}")
        report.append(f"Consistency ratio CR: {self.results['CR']:.4f}")
        report.append(f"Average correlation: {self.results['mean_correlation']:.4f}")
        report.append(f"Chi-square statistic: {self.results['chi2_statistic']:.4f}")
        report.append(f"p-value for Chi-square: {self.results['chi2_p_value']:.4f}")
        report.append(f"Degrees of freedom: {self.results['chi2_df']}")

        # Consistency assessment
        cr_threshold = 0.1
        chi2_threshold = 0.05
        if self.results['CR'] < cr_threshold and self.results['chi2_p_value'] > chi2_threshold:
            report.append("Assessment: The matrix is considered consistent.")
        else:
            report.append("Assessment: The matrix has inconsistency.")
        report.append("\n")

        # Inconsistent pairs
        report.append("4. INCONSISTENT PAIRS")
        report.append("-" * 40)

        if self.results['inconsistent_pairs']:
            report.append(f"{'Alternative triplet':<30} {'Inconsistency':<15} {'Recommended value'}")
            for a, b, c, inc in self.results['inconsistent_pairs'][:5]:
                triplet = f"{names[a]} → {names[b]} → {names[c]}"
                recommended = self.current_matrix[a, b] * self.current_matrix[b, c]
                report.append(f"{triplet:<30} {inc:<15.4f} {recommended:<20.2f}")
        else:
            report.append("No inconsistent pairs found.")

        report.append("\n")

        # Sensitivity analysis
        report.append("5. SENSITIVITY ANALYSIS")
        report.append("-" * 40)

        if hasattr(self, 'sensitivity_results') and self.sensitivity_results:
            i, j = self.sensitivity_results['most_sensitive_pair']
            report.append(f"Most sensitive pair: {names[i]} vs {names[j]}")
            report.append(f"Maximum sensitivity: {self.sensitivity_results['max_sensitivity']:.6f}")
            report.append(f"Perturbation factor: {self.sensitivity_results['perturbation_factor'] * 100:.1f}%")
        else:
            report.append("Sensitivity analysis not performed.")

        report.append("\n")

        # Conclusions
        report.append("6. CONCLUSIONS")
        report.append("-" * 40)

        best_idx_eig = np.argmax(self.results['weights_eigenvector'])
        best_idx_line = np.argmax(self.results['weights_line_method'])
        best_idx_ahp = np.argmax(weights_ahp)

        report.append(f"By the eigenvector method, the best alternative is: {names[best_idx_eig]}")
        report.append(f"By the line method, the best alternative is: {names[best_idx_line]}")
        report.append(f"By the AHP method, the best alternative is: {names[best_idx_ahp]}")

        if best_idx_eig == best_idx_line == best_idx_ahp:
            report.append("Different methods give the same result.")
        else:
            report.append("Different methods give different results.")

        cr_threshold = 0.1
        if self.results['CR'] < cr_threshold:
            report.append("The matrix is consistent (CR < 0.1).")
        else:
            report.append("The matrix is not fully consistent (CR >= 0.1).")

        report.append("\n")
        report.append("Author: Stanislav Kolosov")
        report.append("Date: " + datetime.now().strftime("%Y-%m-%d"))

        return "\n".join(report)

    def export_to_docx(self, report, file_path):
        """Export report to Word document."""
        try:
            doc = Document()

            # Add title
            title = doc.add_heading('REPORT ON LABORATORY WORK #7', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add subtitle
            subtitle = doc.add_paragraph('Decision Support and Improving Consistency of Expert Evaluations')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.bold = True

            # Add variant information
            variant_info = doc.add_paragraph('Variant #1: Information Systems Development')
            variant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            variant_info.runs[0].font.bold = True

            # Add date
            date_info = doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add report sections
            sections = report.split('\n\n')
            for section in sections:
                if section.strip():
                    # If the section starts with a number and a dot, it's a heading
                    if section.strip()[0].isdigit() and '.' in section.strip():
                        doc.add_heading(section.strip(), level=1)
                    else:
                        # Otherwise, add as a regular paragraph
                        doc.add_paragraph(section.strip())

            # Save document
            doc.save(file_path)
        except Exception as e:
            raise e

    def export_to_pdf(self, report, file_path):
        """Export report to PDF."""
        try:
            # Create PDF document
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Add title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(1 * inch, height - 1 * inch, "REPORT ON LABORATORY WORK #7")

            # Add subtitle
            c.setFont("Helvetica-Bold", 12)
            c.drawString(1 * inch, height - 1.5 * inch, "Decision Support and Improving Consistency of Expert Evaluations")

            # Add variant information
            c.setFont("Helvetica-Bold", 12)
            c.drawString(1 * inch, height - 1.8 * inch, "Variant #1: Information Systems Development")

            # Add date
            c.setFont("Helvetica", 10)
            c.drawString(1 * inch, height - 2.1 * inch, f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # Add report content
            c.setFont("Helvetica", 10)
            y_position = height - 2.5 * inch

            # Split report into lines
            lines = report.split('\n')

            for line in lines:
                if y_position < 1 * inch:
                    c.showPage()
                    y_position = height - 1 * inch

                c.drawString(1 * inch, y_position, line)
                y_position -= 12

            # Save PDF
            c.save()
        except Exception as e:
            raise e

    def export_results(self):
        """Export analysis results to a file."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Export results", "First, perform matrix analysis.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[
                ("Excel files", "*.xlsx"),
                ("CSV files", "*.csv"),
                ("Text files", "*.txt"),
                ("LaTeX files", "*.tex")
            ],
            title="Export results"
        )

        if not file_path:
            return

        try:
            # Create DataFrame with results
            size = self.current_matrix.shape[0]
            names = [var.get() for var in self.alternative_names]

            # Weights table
            weights_data = {
                "Alternative": names,
                "Eigenvector method": self.results['weights_eigenvector'],
                "Log. least sq. method": self.results['weights_log_least_squares'],
                "Geometric mean method": self.results['weights_geometric_mean'],
                "Line method": self.results['weights_line_method'],
                "AHP method": calculate_weights_ahp(self.current_matrix)[0]
            }
            weights_df = pd.DataFrame(weights_data)

            # Consistency indicators
            consistency_data = {
                "Metric": [
                    "Eigenvalue λ_max",
                    "Consistency index CI",
                    "Consistency ratio CR",
                    "Average correlation",
                    "Chi-square statistic",
                    "p-value for Chi-square",
                    "Degrees of freedom",
                    "Matrix size"
                ],
                "Value": [
                    f"{self.results['lambda_max']:.4f}", f"{self.results['CI']:.4f}", f"{self.results['CR']:.4f}",
                    f"{self.results['mean_correlation']:.4f}", f"{self.results['chi2_statistic']:.4f}",
                    f"{self.results['chi2_p_value']:.4f}", f"{self.results['chi2_df']}",
                    f"{size}x{size}"
                ]
            }
            consistency_df = pd.DataFrame(consistency_data)

            # Inconsistent pairs
            inconsistent_data = []

            for i, (a, b, c, inc) in enumerate(self.results['inconsistent_pairs']):
                inconsistent_data.append({
                    "Alternative triplet": f"{names[a]} → {names[b]} → {names[c]}",
                    "Inconsistency degree": f"{inc:.4f}",
                    "Recommended value": f"{self.current_matrix[a, b] * self.current_matrix[b, c]:.2f}"
                })

            inconsistent_df = pd.DataFrame(inconsistent_data)

            # Original matrix
            matrix_df = pd.DataFrame(self.current_matrix, index=names, columns=names)

            # Save depending on extension
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.xlsx':
                with pd.ExcelWriter(file_path) as writer:
                    weights_df.to_excel(writer, sheet_name='Weights', index=False)
                    consistency_df.to_excel(writer, sheet_name='Consistency', index=False)
                    inconsistent_df.to_excel(writer, sheet_name='Inconsistent pairs', index=False)
                    matrix_df.to_excel(writer, sheet_name='Matrix')
            elif ext == '.csv':
                # For CSV, save only main results
                output = io.StringIO()
                weights_df.to_csv(output, index=False)

                output.write("\n\nConsistency indicators:\n")
                consistency_df.to_csv(output, index=False)

                output.write("\n\nInconsistent pairs:\n")
                inconsistent_df.to_csv(output, index=False)

                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(output.getvalue())
            elif ext == '.tex':
                # Export to LaTeX
                latex_weights = export_to_latex_table(weights_df, "Alternative weights", "tab:weights")
                latex_consistency = export_to_latex_table(consistency_df, "Consistency indicators", "tab:consistency")

                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("\\documentclass{article}\n")
                    f.write("\\usepackage[utf8]{inputenc}\n")
                    f.write("\\usepackage[russian]{babel}\n")
                    f.write("\\usepackage{amsmath}\n")
                    f.write("\\usepackage{graphicx}\n")
                    f.write("\\usepackage{array}\n")
                    f.write("\\begin{document}\n")
                    f.write("\\title{Report on Laboratory Work #7}\n")
                    f.write("\\author{Stanislav Kolosov}\n")
                    f.write("\\date{" + datetime.now().strftime("%Y-%m-%d") + "}\n")
                    f.write("\\maketitle\n")
                    f.write("\\section{Alternative weights}\n")
                    f.write(latex_weights)
                    f.write("\\section{Consistency indicators}\n")
                    f.write(latex_consistency)
                    f.write("\\end{document}\n")
            else: # .txt
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("===== ALTERNATIVE WEIGHTS =====\n")
                    for _, row in weights_df.iterrows():
                        f.write(f"{row['Alternative']}: "
                                f"Eigenvector={row['Eigenvector method']:.6f}, "
                                f"Log. least sq.={row['Log. least sq. method']:.6f}, "
                                f"Geometric mean={row['Geometric mean method']:.6f}, "
                                f"Line method={row['Line method']:.6f}, "
                                f"AHP={row['AHP method']:.6f}\n")

                    f.write("\n===== CONSISTENCY INDICATORS =====\n")
                    for _, row in consistency_df.iterrows():
                        f.write(f"{row['Metric']}: {row['Value']}\n")

                    f.write("\n===== INCONSISTENT PAIRS =====\n")
                    for _, row in inconsistent_df.iterrows():
                        f.write(
                            f"{row['Alternative triplet']}: {row['Inconsistency degree']} (recommended: {row['Recommended value']})\n"
                        )

                    f.write("\n===== ORIGINAL MATRIX =======\n")
                    for i, name in enumerate(names):
                        f.write(f"{name}: " + ", ".join([f"{self.current_matrix[i, j]:.2f}" for j in range(size)]) + "\n")

            self.log_message(f"Results exported to file: {os.path.basename(file_path)}")
            messagebox.showinfo("Export results", f"Results successfully saved to file:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Export error", f"Failed to export results: {str(e)}")
            self.log_message(f"Error exporting results: {str(e)}")

    def save_session(self):
        """Save the analysis session."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Save session", "First, perform matrix analysis.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            title="Save session"
        )

        if not file_path:
            return

        try:
            # Prepare session data
            session_data = {
                "timestamp": datetime.now().isoformat(),
                "matrix_size": self.matrix_size.get(),
                "alternative_names": [var.get() for var in self.alternative_names],
                "matrix": self.current_matrix.tolist(),
                "results": {
                    "weights_eigenvector": self.results['weights_eigenvector'].tolist(),
                    "weights_log_least_squares": self.results['weights_log_least_squares'].tolist(),
                    "weights_geometric_mean": self.results['weights_geometric_mean'].tolist(),
                    "weights_line_method": self.results['weights_line_method'].tolist(),
                    "lambda_max": float(self.results['lambda_max']),
                    "CI": float(self.results['CI']),
                    "CR": float(self.results['CR']),
                    "mean_correlation": float(self.results['mean_correlation']),
                    "chi2_statistic": float(self.results['chi2_statistic']),
                    "chi2_p_value": float(self.results['chi2_p_value']),
                    "chi2_df": int(self.results['chi2_df']),
                    "inconsistent_pairs": [[int(a), int(b), int(c), float(inc)] for a, b, c, inc in self.results['inconsistent_pairs']],
                    "consistent_matrix": self.results['consistent_matrix'].tolist()
                },
                "sensitivity_results": None
            }

            if hasattr(self, 'sensitivity_results') and self.sensitivity_results:
                session_data["sensitivity_results"] = {
                    "max_sensitivity": float(self.sensitivity_results['max_sensitivity']),
                    "most_sensitive_pair": [int(i) for i in self.sensitivity_results['most_sensitive_pair']],
                    "perturbation_factor": float(self.sensitivity_results['perturbation_factor'])
                }

            # Save session
            success = save_session(session_data, file_path)

            if success:
                self.session_file = file_path
                self.log_message(f"Session saved to file: {os.path.basename(file_path)}")
                messagebox.showinfo("Save session", f"Session successfully saved to file:\n{file_path}")
            else:
                raise Exception("Failed to save session")
        except Exception as e:
            messagebox.showerror("Save error", f"Failed to save session: {str(e)}")
            self.log_message(f"Error saving session: {str(e)}")

    def load_session(self):
        """Load an analysis session."""
        file_path = filedialog.askopenfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            title="Load session"
        )

        if not file_path:
            return

        try:
            # Load session
            session_data = load_session(file_path)
            if session_data is None:
                raise Exception("Failed to load session")

            # Restore application state
            self.matrix_size.set(session_data["matrix_size"])
            self.create_matrix_entries(session_data["matrix_size"])

            # Set alternative names
            for i, name in enumerate(session_data["alternative_names"]):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Restore matrix
            matrix_data = np.array(session_data["matrix"])
            self.current_matrix = matrix_data

            # Restore results
            results_data = session_data["results"]
            self.results = {
                'weights_eigenvector': np.array(results_data["weights_eigenvector"]),
                'weights_log_least_squares': np.array(results_data["weights_log_least_squares"]),
                'weights_geometric_mean': np.array(results_data["weights_geometric_mean"]),
                'weights_line_method': np.array(results_data["weights_line_method"]),
                'lambda_max': results_data["lambda_max"],
                'CI': results_data["CI"],
                'CR': results_data["CR"],
                'mean_correlation': results_data["mean_correlation"],
                'chi2_statistic': results_data["chi2_statistic"],
                'chi2_p_value': results_data["chi2_p_value"],
                'chi2_df': results_data["chi2_df"],
                'inconsistent_pairs': [(a, b, c, inc) for a, b, c, inc in results_data["inconsistent_pairs"]],
                'consistent_matrix': np.array(results_data["consistent_matrix"])
            }

            # Restore sensitivity analysis if available
            if "sensitivity_results" in session_data and session_data["sensitivity_results"]:
                self.sensitivity_results = {
                    'max_sensitivity': session_data["sensitivity_results"]["max_sensitivity"],
                    'most_sensitive_pair': tuple(session_data["sensitivity_results"]["most_sensitive_pair"]),
                    'perturbation_factor': session_data["sensitivity_results"]["perturbation_factor"]
                }

            # Fill the matrix in the interface
            for i, row in enumerate(matrix_data):
                for j, val in enumerate(row):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{val:.2f}")

            # Update all tabs
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.session_file = file_path

            self.log_message(f"Session loaded from file: {os.path.basename(file_path)}")
            messagebox.showinfo("Load session", f"Session successfully loaded from file:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Load error", f"Failed to load session: {str(e)}")
            self.log_message(f"Error loading session: {str(e)}")

    def import_from_csv(self):
        """Import matrix from CSV file."""
        file_path = filedialog.askopenfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            title="Import from CSV"
        )

        if not file_path:
            return

        try:
            # Load data from CSV
            df = pd.read_csv(file_path, header=None)
            # Check dimensions
            if df.shape[0] != df.shape[1]:
                raise ValueError("Matrix must be square")
            # Set matrix size
            size = df.shape[0]
            self.matrix_size.set(size)
            self.create_matrix_entries(size)
            # Set alternative names (if present in CSV)
            if len(df.columns) > size:
                # First column contains row names
                row_names = df.iloc[:, 0].astype(str).tolist()
                col_names = df.iloc[0, :].astype(str).tolist()

                data_start_row = 1
                data_start_col = 1
            else:
                # Only data, no names
                row_names = [f"Alternative {i + 1}" for i in range(size)]
                col_names = [f"Alternative {j + 1}" for j in range(size)]
                data_start_row = 0
                data_start_col = 0

            # Set alternative names
            for i, name in enumerate(row_names):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Fill matrix with data
            matrix_data = df.iloc[data_start_row: data_start_row + size, data_start_col: data_start_col + size].values.astype(float)

            # Check that the matrix is positive definite
            for i in range(size):
                for j in range(size):
                    if i == j:
                        if matrix_data[i, j] != 1:
                            raise ValueError(f"Diagonal element [{i}, {j}] must be 1, found: {matrix_data[i, j]}")
                    else:
                        if matrix_data[i, j] <= 0:
                            raise ValueError(f"Element [{i}, {j}] must be positive, found: {matrix_data[i, j]}")
                        if abs(matrix_data[i, j] * matrix_data[j, i] - 1) > 1e-10:
                            raise ValueError(f"Elements [{i}, {j}] and [{j}, {i}] are not inverses of each other")

            # Set matrix
            self.current_matrix = matrix_data

            # Fill interface
            for i, row in enumerate(matrix_data):
                for j, val in enumerate(row):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{val:.2f}")

            # Perform analysis
            self.results = analyze_consistency(matrix_data)
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.log_message(f"Matrix successfully imported from CSV file: {os.path.basename(file_path)}")
            messagebox.showinfo("Import from CSV", f"Matrix successfully imported from file:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Import error", f"Failed to import matrix from CSV: {str(e)}")
            self.log_message(f"Error importing from CSV: {str(e)}")

    def import_from_excel(self):
        """Import matrix from Excel file."""
        file_path = filedialog.askopenfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            title="Import from Excel"
        )

        if not file_path:
            return

        try:
            # Load data from Excel
            excel_file = pd.ExcelFile(file_path)

            # Prompt to select a sheet
            if len(excel_file.sheet_names) > 1:
                sheet_name = simpledialog.askstring("Select sheet", f"Select a sheet:\n{', '.join(excel_file.sheet_names)}", parent=self.root)
                if sheet_name not in excel_file.sheet_names:
                    sheet_name = excel_file.sheet_names[0]
            else:
                sheet_name = excel_file.sheet_names[0]

            # Load the selected sheet
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

            # Check dimensions
            if df.shape[0] != df.shape[1]:
                raise ValueError("Matrix must be square")

            # Set matrix size
            size = df.shape[0]
            self.matrix_size.set(size)
            self.create_matrix_entries(size)

            # Set alternative names (if present in Excel)
            if len(df.columns) > size:
                # First column contains row names
                row_names = df.iloc[:, 0].astype(str).tolist()
                col_names = df.iloc[0, :].astype(str).tolist()
                data_start_row = 1
                data_start_col = 1
            else:
                # Only data, no names
                row_names = [f"Alternative {i + 1}" for i in range(size)]
                col_names = [f"Alternative {j + 1}" for j in range(size)]
                data_start_row = 0
                data_start_col = 0

            # Set alternative names
            for i, name in enumerate(row_names):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Fill matrix with data
            matrix_data = df.iloc[data_start_row: data_start_row + size, data_start_col: data_start_col + size].values.astype(float)

            # Check that the matrix is positive definite
            for i in range(size):
                for j in range(size):
                    if i == j:
                        if matrix_data[i, j] != 1:
                            raise ValueError(f"Diagonal element [{i}, {j}] must be 1, found: {matrix_data[i, j]}")
                    else:
                        if matrix_data[i, j] <= 0:
                            raise ValueError(f"Element [{i}, {j}] must be positive, found: {matrix_data[i, j]}")
                        if abs(matrix_data[i, j] * matrix_data[j, i] - 1) > 1e-10:
                            raise ValueError(f"Elements [{i}, {j}] and [{j}, {i}] are not inverses of each other")

            # Set matrix
            self.current_matrix = matrix_data

            # Fill interface
            for i in range(size):
                for j in range(size):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{matrix_data[i, j]:.2f}")

            # Perform analysis
            self.results = analyze_consistency(matrix_data)
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.log_message(f"Matrix successfully imported from Excel file: {os.path.basename(file_path)}")
            messagebox.showinfo("Import from Excel", f"Matrix successfully imported from file:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Import error", f"Failed to import matrix from Excel: {str(e)}")
            self.log_message(f"Error importing from Excel: {str(e)}")

    def show_about(self):
        """Show information about the program."""
        about_text = (
            "Decision Support and Improving Consistency of Expert Evaluations Program\n"
            "Laboratory Work #7 on the subject 'Fundamentals of AI'\n"
            "Variant 1: Information Systems Development\n\n"
            "Main functions:\n"
            "- Analysis of pairwise comparison matrix\n"
            "- Calculation of alternative weights using four methods\n"
            "- Evaluation of consistency of expert evaluations\n"
            "- Identification of the most inconsistent pairwise comparisons\n"
            "- Sensitivity analysis and inconsistency correction\n"
            "- 3D visualization of results\n"
            "- Export of full report\n\n"
            "Developed for educational purposes\n"
            "Author: Stanislav Kolosov"
        )
        messagebox.showinfo("About", about_text)

    def toggle_theme(self):
        """Toggle the interface color theme."""
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.root.config(bg=self.themes["dark"]["bg"])
            self.style.configure("TFrame", background=self.themes["dark"]["bg"])
            self.style.configure("TLabel", background=self.themes["dark"]["bg"], foreground=self.themes["dark"]["fg"])
            self.style.configure("Header.TLabel", background=self.themes["dark"]["bg"], foreground=self.themes["dark"]["fg"])
        else:
            self.current_theme = "light"
            self.root.config(bg=self.themes["light"]["bg"])
            self.style.configure("TFrame", background=self.themes["light"]["bg"])
            self.style.configure("TLabel", background=self.themes["light"]["bg"], foreground=self.themes["light"]["fg"])
            self.style.configure("Header.TLabel", background=self.themes["light"]["bg"], foreground=self.themes["light"]["fg"])

        self.log_message(f"Interface theme changed to {self.current_theme}.")

if __name__ == "__main__":
    root = tk.Tk()
    app = DecisionSupportApp(root)
    root.mainloop()
