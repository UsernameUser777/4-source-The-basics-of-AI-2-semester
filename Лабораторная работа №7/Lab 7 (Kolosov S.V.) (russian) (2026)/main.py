# -*- coding: utf-8 -*-

"""
Лабораторная работа №7: Поддержка принятия решений и повышение согласованности экспертных оценок
Вариант №1

Программа для анализа матриц парных сравнений, вычисления весов альтернатив и оценки согласованности экспертных оценок.
"""

import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext, colorchooser, simpledialog
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from mpl_toolkits.mplot3d import Axes3D
import pandas as pd
import os
import io
import json
from decimal import Decimal
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import sys
from datetime import datetime
from decision_support import *

class DecisionSupportApp:
    """Класс для графического интерфейса приложения поддержки принятия решений."""

    def __init__(self, root):
        """Инициализация приложения."""
        self.root = root
        self.root.title("Поддержка принятия решений (Вариант 1)")
        self.root.geometry("1400x900")

        # Настройка стилей
        self.style = ttk.Style()
        self.style.configure("TFrame", background="#f0f0f0")
        self.style.configure("TButton", padding=6)
        self.style.configure("TLabel", background="#f0f0f0", font=("Arial", 10))
        self.style.configure("Header.TLabel", font=("Arial", 12, "bold"))

        # Цветовые темы
        self.themes = {
            "light": {"bg": "#f0f0f0", "fg": "black", "highlight": "blue"},
            "dark": {"bg": "#2b2b2b", "fg": "white", "highlight": "cyan"}
        }
        self.current_theme = "light"

        # Переменные состояния
        self.matrix_size = tk.IntVar(value=4)
        self.alternative_names = []
        self.matrix_entries = []
        self.current_matrix = None
        self.results = None
        self.sensitivity_results = None
        self.session_file = None

        # Добавленные цвета для визуализации
        self.color_map = plt.cm.Set3
        self.highlight_color = 'red'

        # Журнал операций для отката
        self.operation_history = []
        self.max_history_size = 50

        # Создание интерфейса
        self.create_widgets()

        # Загрузка примера данных для варианта 1
        self.root.after(100, self.load_example_data)

    def create_widgets(self):
        """Создание элементов интерфейса."""
        # Создаем основные фреймы
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Левая панель для ввода матрицы
        left_frame = ttk.LabelFrame(main_frame, text="Матрица парных сравнений")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Правая панель для результатов
        right_frame = ttk.Frame(main_frame)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Панель управления
        control_frame = ttk.LabelFrame(right_frame, text="Управление")
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        # Панель результатов
        results_frame = ttk.LabelFrame(right_frame, text="Результаты анализа")
        results_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=(0, 5))

        # --- Левая панель: матрица парных сравнений ---
        # Фрейм для настроек матрицы
        matrix_settings_frame = ttk.Frame(left_frame)
        matrix_settings_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(matrix_settings_frame, text="Размер матрицы:").pack(side=tk.LEFT, padx=(0, 5))
        size_spinbox = ttk.Spinbox(matrix_settings_frame, from_=2, to=10,
                                   textvariable=self.matrix_size, width=5,
                                   command=self.update_matrix_size)
        size_spinbox.pack(side=tk.LEFT, padx=(0, 10))

        ttk.Button(matrix_settings_frame, text="Обновить матрицу",
                   command=self.update_matrix_size).pack(side=tk.LEFT)

        # Фрейм для имен альтернатив
        names_frame = ttk.Frame(left_frame)
        names_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(names_frame, text="Имена альтернатив:").pack(anchor=tk.W, padx=5, pady=(0, 5))

        self.names_entries_frame = ttk.Frame(names_frame)
        self.names_entries_frame.pack(fill=tk.X)

        # Фрейм для самой матрицы
        self.matrix_frame = ttk.Frame(left_frame)
        self.matrix_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем матрицу по умолчанию (4x4)
        self.create_matrix_entries(4)

        # --- Панель управления ---
        # Кнопки управления
        btn_frame = ttk.Frame(control_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(btn_frame, text="Анализировать",
                   command=self.analyze_matrix, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Загрузить пример",
                   command=self.load_example_data, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="Экспорт результатов",
                   command=self.export_results, width=15).pack(side=tk.LEFT, padx=2)
        ttk.Button(btn_frame, text="О программе",
                   command=self.show_about, width=10).pack(side=tk.RIGHT, padx=2)

        # Кнопки дополнительных функций
        extra_btn_frame = ttk.Frame(control_frame)
        extra_btn_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(extra_btn_frame, text="Анализ чувствительности",
                   command=self.analyze_sensitivity, width=20).pack(side=tk.LEFT, padx=2)
        ttk.Button(extra_btn_frame, text="Корректировать несогласованность",
                   command=self.adjust_inconsistencies, width=22).pack(side=tk.LEFT, padx=2)
        ttk.Button(extra_btn_frame, text="Сгенерировать согласованную матрицу",
                   command=self.generate_consistent_matrix, width=28).pack(side=tk.LEFT, padx=2)

        # Кнопки импорта/экспорта
        import_export_frame = ttk.Frame(control_frame)
        import_export_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(import_export_frame, text="Сохранить сессию",
                   command=self.save_session).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Загрузить сессию",
                   command=self.load_session).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Импорт из CSV",
                   command=self.import_from_csv).pack(side=tk.LEFT, padx=2)
        ttk.Button(import_export_frame, text="Импорт из Excel",
                   command=self.import_from_excel).pack(side=tk.LEFT, padx=2)

        # Кнопка для переключения тем
        ttk.Button(import_export_frame, text="Переключить тему",
                   command=self.toggle_theme).pack(side=tk.RIGHT, padx=2)

        # --- Панель результатов ---
        # Вкладки для результатов
        self.results_notebook = ttk.Notebook(results_frame)
        self.results_notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Вкладка: веса альтернатив
        self.weights_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.weights_frame, text="Веса альтернатив")
        self.create_weights_tab()

        # Вкладка: сравнение методов
        self.comparison_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.comparison_frame, text="Сравнение методов")
        self.create_comparison_tab()

        # Вкладка: согласованность
        self.consistency_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.consistency_frame, text="Согласованность")
        self.create_consistency_tab()

        # Вкладка: несогласованные пары
        self.inconsistent_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.inconsistent_frame, text="Несогласованные пары")
        self.create_inconsistent_pairs_tab()

        # Вкладка: 3D визуализация
        self.visualization_3d_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.visualization_3d_frame, text="3D Визуализация")
        self.create_3d_visualization_tab()

        # Вкладка: лог выполнения
        self.log_frame = ttk.Frame(self.results_notebook)
        self.results_notebook.add(self.log_frame, text="Лог выполнения")
        self.create_log_tab()

        # Добавляем контекстную справку ко всем элементам
        self.add_context_help()

    def add_context_help(self):
        """Добавление контекстной справки ко всем элементам интерфейса."""
        help_texts = {
            "matrix_size": "Размер матрицы парных сравнений (от 2 до 10)",
            "alternative_names": "Названия сравниваемых альтернатив",
            "analyze_button": "Выполнить анализ матрицы и вычислить веса альтернатив",
            "sensitivity_button": "Проанализировать чувствительность результатов к изменениям в матрице",
            "adjust_button": "Автоматически скорректировать наиболее несогласованные парные сравнения",
            "save_session": "Сохранить всю текущую сессию анализа (матрицу, настройки, результаты)",
            "load_session": "Загрузить ранее сохраненную сессию анализа",
            "weights_tab": "Просмотр весов альтернатив, рассчитанных различными методами",
            "comparison_tab": "Сравнение результатов, полученных разными методами",
            "consistency_tab": "Показатели согласованности экспертных оценок",
            "inconsistent_tab": "Список наиболее несогласованных парных сравнений",
            "3d_tab": "3D визуализация матрицы парных сравнений и весов альтернатив",
            "log_tab": "Журнал выполненных операций и сообщений системы"
        }

    def show_tooltip(self, text):
        """Отображение всплывающей подсказки."""
        try:
            self.tooltip = tk.Toplevel(self.root)
            self.tooltip.wm_overrideredirect(True)
            self.tooltip.wm_geometry(f"+{self.root.winfo_pointerx() + 10}+{self.root.winfo_pointery() + 10}")

            label = ttk.Label(self.tooltip, text=text, background="lightyellow", relief="solid", borderwidth=1)
            label.pack()
        except Exception:
            pass

    def hide_tooltip(self):
        """Скрытие всплывающей подсказки."""
        try:
            if hasattr(self, 'tooltip'):
                self.tooltip.destroy()
        except Exception:
            pass

    def create_matrix_entries(self, size):
        """Создание полей для ввода матрицы заданного размера."""
        # Очищаем предыдущую матрицу
        for widget in self.matrix_frame.winfo_children():
            widget.destroy()
        self.matrix_entries = []
        self.alternative_names = []

        # Создаем заголовки столбцов
        for j in range(size):
            label = ttk.Label(self.matrix_frame, text=f"Альтернатива {j + 1}", font=("Arial", 9, "bold"))
            label.grid(row=0, column=j + 1, padx=2, pady=2)

        # Создаем поля ввода для матрицы
        for i in range(size):
            # Заголовок строки
            label = ttk.Label(self.matrix_frame, text=f"Альтернатива {i + 1}", font=("Arial", 9, "bold"))
            label.grid(row=i + 1, column=0, padx=2, pady=2)

            row_entries = []
            for j in range(size):
                entry_var = tk.StringVar(value="1.0" if i == j else "3.0" if i < j else "0.33")
                entry = ttk.Entry(self.matrix_frame, textvariable=entry_var, width=8)

                # Диагональные элементы не редактируются (должны быть 1)
                if i == j:
                    entry.config(state="readonly")

                # Добавляем проверку ввода
                entry.bind("<FocusOut>", lambda event, r=i, c=j: self.validate_entry(r, c))
                entry.bind("<Return>", lambda event, r=i, c=j: self.on_enter_pressed(r, c))

                entry.grid(row=i + 1, column=j + 1, padx=2, pady=2)
                row_entries.append(entry_var)
            self.matrix_entries.append(row_entries)

        # Обновляем имена альтернатив
        self.update_alternative_names(size)

    def validate_entry(self, i, j):
        """Проверка корректности введенного значения."""
        if i == j:
            return

        try:
            value = float(self.matrix_entries[i][j].get())
            if value <= 0:
                raise ValueError("Значение должно быть положительным")

            # Проверяем транзитивность
            message = f"Элемент [{i + 1},{j + 1}] установлен в {value:.2f}"
            self.log_message(message)

        except ValueError as e:
            messagebox.showerror("Ошибка ввода", f"Некорректное значение: {str(e)}")
            self.log_message(f"Ошибка ввода в элемент [{i + 1},{j + 1}]")
            self.matrix_entries[i][j].set("1.0")

    def on_enter_pressed(self, i, j):
        """Действие при нажатии Enter."""
        self.log_message(f"Изменено значение в [{i + 1},{j + 1}]")
        # Добавляем в историю операций
        self.add_to_operation_history(f"Изменение элемента [{i + 1},{j + 1}]")

    def update_alternative_names(self, size):
        """Обновление полей для имен альтернатив."""
        # Очищаем предыдущие поля
        for widget in self.names_entries_frame.winfo_children():
            widget.destroy()
        self.alternative_names = []

        for i in range(size):
            frame = ttk.Frame(self.names_entries_frame)
            frame.pack(fill=tk.X, pady=2)

            ttk.Label(frame, text=f"Альтернатива {i + 1}:").pack(side=tk.LEFT, padx=(0, 5))
            name_var = tk.StringVar(value=f"Вариант {i + 1}")
            entry = ttk.Entry(frame, textvariable=name_var, width=20)
            entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
            self.alternative_names.append(name_var)

    def update_matrix_size(self):
        """Обновление размера матрицы."""
        size = self.matrix_size.get()
        self.create_matrix_entries(size)
        self.update_alternative_names(size)

    def get_matrix_from_entries(self):
        """Получение матрицы из полей ввода."""
        size = len(self.matrix_entries)
        matrix = np.ones((size, size))

        try:
            for i in range(size):
                for j in range(size):
                    if i != j:
                        val = float(self.matrix_entries[i][j].get())
                        # Проверка на положительность
                        if val <= 0:
                            raise ValueError("Значения матрицы должны быть положительными")
                        matrix[i, j] = val
                        # Для симметричных элементов (j, i) должно быть 1/val
                        matrix[j, i] = 1 / val
            return matrix
        except ValueError as e:
            messagebox.showerror("Ошибка ввода", f"Некорректное значение в матрице: {str(e)}")
            return None

    def load_example_data(self):
        """Загрузка примера данных для варианта 1."""
        # Для варианта 1 используем пример с выбором места для строительства объекта
        size = 4
        self.matrix_size.set(size)
        self.create_matrix_entries(size)
        self.update_alternative_names(size)

        # Устанавливаем имена альтернатив согласно варианту №1
        alternative_names = ["Ведение проекта", "Анализ требований", "Проектирование", "Реализация"]
        for i, name in enumerate(alternative_names):
            if i < len(self.alternative_names):
                self.alternative_names[i].set(name)

        # Заполняем матрицу примером
        example_matrix = [
            [1, 3, 5, 7],
            [1 / 3, 1, 3, 5],
            [1 / 5, 1 / 3, 1, 3],
            [1 / 7, 1 / 5, 1 / 3, 1]
        ]

        for i in range(size):
            for j in range(size):
                if i != j:
                    self.matrix_entries[i][j].set(f"{example_matrix[i][j]:.2f}")

        self.log_message("Загружен пример данных для варианта 1: разработка информационных систем.")
        self.log_message("Матрица парных сравнений заполнена примерными значениями.")

        # Выполняем анализ после загрузки примера
        self.analyze_matrix()

    def analyze_matrix(self):
        """Анализ матрицы парных сравнений."""
        matrix = self.get_matrix_from_entries()
        if matrix is None:
            return

        self.current_matrix = matrix

        try:
            # Выполняем анализ согласованности
            self.results = analyze_consistency(matrix)

            # Обновляем все вкладки
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.update_3d_visualization()

            self.log_message("Анализ матрицы успешно выполнен.")

        except Exception as e:
            messagebox.showerror("Ошибка анализа", f"Произошла ошибка при анализе матрицы: {str(e)}")
            self.log_message(f"Ошибка: {str(e)}")

    def analyze_sensitivity(self):
        """Анализ чувствительности матрицы."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Анализ чувствительности", "Сначала выполните анализ матрицы.")
            return

        try:
            # Выполняем анализ чувствительности
            self.sensitivity_results = create_sensitivity_analysis(
                self.current_matrix, self.results, perturbation_factor=0.1
            )

            # Отображаем результаты
            i, j = self.sensitivity_results['most_sensitive_pair']
            names = [var.get() for var in self.alternative_names]

            message = (
                f"Наиболее чувствительная пара: {names[i]} vs {names[j]}\n"
                f"Максимальная чувствительность: {self.sensitivity_results['max_sensitivity']:.6f}\n"
                f"Фактор возмущения: {self.sensitivity_results['perturbation_factor'] * 100:.1f}%"
            )
            messagebox.showinfo("Анализ чувствительности", message)
            self.log_message(f"Выполнен анализ чувствительности: {message}")

            # Обновляем 3D визуализацию
            self.update_3d_visualization()

        except Exception as e:
            messagebox.showerror("Ошибка анализа", f"Произошла ошибка при анализе чувствительности: {str(e)}")
            self.log_message(f"Ошибка анализа чувствительности: {str(e)}")

    def adjust_inconsistencies(self):
        """Корректировка несогласованных парных сравнений."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Корректировка несогласованности", "Сначала выполните анализ матрицы.")
            return

        try:
            # Выполняем корректировку
            corrected_matrix, changes = adjust_inconsistent_pairs(
                self.current_matrix, self.results['inconsistent_pairs'], threshold=0.1
            )

            if not changes:
                messagebox.showinfo("Корректировка", "Несогласованность находится в допустимых пределах.")
                self.log_message("Корректировка несогласованности: изменения не требуются.")
                return

            # Предлагаем пользователю применить изменения
            confirm = messagebox.askyesno(
                "Корректировка",
                f"Обнаружено {len(changes)} несогласованных пар.\nПрименить автоматические исправления?"
            )

            if confirm:
                # Применяем изменения к интерфейсу
                for change in changes:
                    i, j = change['pair']
                    self.matrix_entries[i][j].set(f"{change['corrected']:.2f}")
                    self.matrix_entries[j][i].set(f"{1 / change['corrected']:.2f}")
                # Обновляем матрицу и перезапускаем анализ
                self.current_matrix = corrected_matrix
                self.results = analyze_consistency(corrected_matrix)

                # Обновляем все вкладки
                self.update_weights_tab()
                self.update_comparison_tab()
                self.update_consistency_tab()
                self.update_inconsistent_pairs_tab()
                self.update_3d_visualization()

                self.log_message(f"Применены автоматические исправления для {len(changes)} несогласованных пар.")
                self.log_message("Матрица была скорректирована и повторно проанализирована.")

        except Exception as e:
            messagebox.showerror("Ошибка корректировки", f"Произошла ошибка при корректировке: {str(e)}")
            self.log_message(f"Ошибка корректировки: {str(e)}")

    def generate_consistent_matrix(self):
        """Генерация согласованной матрицы."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Генерация матрицы", "Сначала выполните анализ матрицы.")
            return

        try:
            # Получаем согласованную матрицу из результатов
            consistent_matrix = self.results['consistent_matrix']

            # Создаем новое окно для отображения согласованной матрицы
            dialog = tk.Toplevel(self.root)
            dialog.title("Согласованная матрица")
            dialog.geometry("600x500")

            # Создаем фрейм для матрицы
            matrix_frame = ttk.Frame(dialog)
            matrix_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            # Создаем таблицу
            columns = [f"Альтернатива {i + 1}" for i in range(consistent_matrix.shape[0])]
            tree = ttk.Treeview(matrix_frame, columns=columns, show="headings", height=10)

            # Настройка заголовков
            names = [var.get() for var in self.alternative_names]
            for i, name in enumerate(names):
                tree.heading(columns[i], text=name)
                tree.column(columns[i], width=100)

            # Добавляем данные
            for i, row in enumerate(consistent_matrix):
                values = [f"{val:.3f}" for val in row]
                tree.insert("", tk.END, values=values)

            # Добавляем скроллбары
            v_scrollbar = ttk.Scrollbar(matrix_frame, orient=tk.VERTICAL, command=tree.yview)
            h_scrollbar = ttk.Scrollbar(matrix_frame, orient=tk.HORIZONTAL, command=tree.xview)
            tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

            tree.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
            v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
            h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

            # Кнопки управления
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(fill=tk.X, padx=10, pady=10)

            ttk.Button(btn_frame, text="Копировать",
                       command=lambda: self.copy_matrix_to_clipboard(consistent_matrix)).pack(side=tk.LEFT)
            ttk.Button(btn_frame, text="Применить",
                       command=lambda: self.apply_consistent_matrix(consistent_matrix, dialog)).pack(side=tk.LEFT, padx=5)
            ttk.Button(btn_frame, text="Закрыть",
                       command=dialog.destroy).pack(side=tk.RIGHT)

            self.log_message("Сгенерирована согласованная матрица на основе текущих весов.")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка при генерации согласованной матрицы: {str(e)}")
            self.log_message(f"Ошибка генерации согласованной матрицы: {str(e)}")

    def copy_matrix_to_clipboard(self, matrix):
        """Копирование матрицы в буфер обмена."""
        try:
            # Преобразуем матрицу в текстовый формат
            text = ""
            names = [var.get() for var in self.alternative_names]

            # Добавляем заголовки
            text += "\t" + "\t".join(names) + "\n"

            # Добавляем строки матрицы
            for i, row in enumerate(matrix):
                text += names[i] + "\t" + "\t".join([f"{val:.3f}" for val in row]) + "\n"

            self.root.clipboard_clear()
            self.root.clipboard_append(text)

            self.log_message("Согласованная матрица скопирована в буфер обмена.")
            messagebox.showinfo("Копирование", "Матрица скопирована в буфер обмена.")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось скопировать матрицу: {str(e)}")
            self.log_message(f"Ошибка копирования матрицы: {str(e)}")

    def apply_consistent_matrix(self, matrix, dialog):
        """Применение согласованной матрицы к основному интерфейсу."""
        try:
            # Применяем значения к основному интерфейсу
            for i in range(matrix.shape[0]):
                for j in range(matrix.shape[1]):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{matrix[i, j]:.3f}")

            # Обновляем основную матрицу
            self.current_matrix = matrix
            self.results = analyze_consistency(matrix)

            # Обновляем все вкладки
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.update_3d_visualization()

            # Закрываем диалоговое окно
            dialog.destroy()

            self.log_message("Согласованная матрица применена к основному интерфейсу.")
            messagebox.showinfo("Применение", "Согласованная матрица успешно применена.")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось применить матрицу: {str(e)}")
            self.log_message(f"Ошибка применения матрицы: {str(e)}")

    def create_weights_tab(self):
        """Создание вкладки для отображения весов альтернатив."""
        # Фрейм для выбора метода
        method_frame = ttk.Frame(self.weights_frame)
        method_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(method_frame, text="Метод вычисления весов:").pack(side=tk.LEFT, padx=(0, 10))
        self.weights_method = tk.StringVar(value="eigenvector")
        methods = [
            ("Метод собственного вектора", "eigenvector"),
            ("Метод логарифмических наименьших квадратов", "log_least_squares"),
            ("Метод среднего геометрического", "geometric_mean"),
            ("Метод \"линия\"", "line_method"),
            ("Метод AHP", "ahp")
        ]

        for text, method in methods:
            rb = ttk.Radiobutton(method_frame, text=text, value=method,
                                 variable=self.weights_method,
                                 command=self.update_weights_display)
            rb.pack(side=tk.LEFT, padx=5)

        # Фрейм для отображения весов
        self.weights_display_frame = ttk.Frame(self.weights_frame)
        self.weights_display_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем график и таблицу
        self.create_weights_display()

    def create_weights_display(self):
        """Создание элементов для отображения весов."""
        # Очищаем предыдущее содержимое
        for widget in self.weights_display_frame.winfo_children():
            widget.destroy()

        # Создаем фрейм для графика
        chart_frame = ttk.Frame(self.weights_display_frame)
        chart_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Создаем график
        self.weights_fig, self.weights_ax = plt.subplots(figsize=(6, 4))
        self.weights_canvas = FigureCanvasTkAgg(self.weights_fig, master=chart_frame)
        self.weights_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Создаем фрейм для таблицы
        table_frame = ttk.Frame(self.weights_display_frame)
        table_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Создаем таблицу
        columns = ("alternative", "weight")
        self.weights_tree = ttk.Treeview(table_frame, columns=columns, show="headings")

        # Настройка заголовков
        self.weights_tree.heading("alternative", text="Альтернатива")
        self.weights_tree.heading("weight", text="Вес")

        # Настройка колонок
        self.weights_tree.column("alternative", width=150)
        self.weights_tree.column("weight", width=100, anchor=tk.E)

        # Добавляем скроллбар
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.weights_tree.yview)
        self.weights_tree.configure(yscrollcommand=scrollbar.set)

        self.weights_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def update_weights_tab(self):
        """Обновление вкладки с весами альтернатив."""
        if self.results is None:
            return

        # Обновляем отображение весов
        self.update_weights_display()

    def update_weights_display(self):
        """Обновление отображения весов в зависимости от выбранного метода."""
        if self.results is None:
            return

        method = self.weights_method.get()

        # Определяем, какие веса использовать
        if method == "eigenvector":
            weights = self.results['weights_eigenvector']
            method_name = "Метод собственного вектора"
        elif method == "log_least_squares":
            weights = self.results['weights_log_least_squares']
            method_name = "Метод логарифмических наименьших квадратов"
        elif method == "geometric_mean":
            weights = self.results['weights_geometric_mean']
            method_name = "Метод среднего геометрического"
        elif method == "ahp":
            weights, _, _, _ = calculate_weights_ahp(self.current_matrix)
            method_name = "Метод AHP"
        else:  # line_method
            weights = self.results['weights_line_method']
            method_name = "Метод \"линия\""

        # Обновляем график
        self.weights_ax.clear()

        # Получаем имена альтернатив
        names = [var.get() for var in self.alternative_names]

        # Строим столбчатую диаграмму
        bars = self.weights_ax.bar(names, weights, color='skyblue')
        self.weights_ax.set_title(f'Веса альтернатив ({method_name})')
        self.weights_ax.set_ylabel('Вес')
        self.weights_ax.set_ylim(0, max(weights) * 1.2)

        # Добавляем значения над столбцами
        for bar in bars:
            height = bar.get_height()
            self.weights_ax.annotate(f'{height:.4f}', xy=(bar.get_x() + bar.get_width() / 2, height),
                                     xytext=(0, 3), textcoords="offset points", ha='center', va='bottom')

        self.weights_ax.tick_params(axis='x', rotation=15)
        self.weights_fig.tight_layout()
        self.weights_canvas.draw()

        # Обновляем таблицу
        for item in self.weights_tree.get_children():
            self.weights_tree.delete(item)

        for i, (name, weight) in enumerate(zip(names, weights)):
            self.weights_tree.insert("", tk.END, values=(name, f"{weight:.6f}"))

        # Определяем лучшую альтернативу
        best_idx = np.argmax(weights)
        best_name = names[best_idx]
        best_weight = weights[best_idx]

        self.log_message(f"Лучшая альтернатива по методу {method_name}: {best_name} (вес = {best_weight:.4f})")

    def create_comparison_tab(self):
        """Создание вкладки для сравнения методов."""
        # Фрейм для отображения сравнения
        self.comparison_display_frame = ttk.Frame(self.comparison_frame)
        self.comparison_display_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем график и таблицу
        self.create_comparison_display()

    def create_comparison_display(self):
        """Создание элементов для отображения сравнения методов."""
        # Очищаем предыдущее содержимое
        for widget in self.comparison_display_frame.winfo_children():
            widget.destroy()

        # Создаем фрейм для графика
        chart_frame = ttk.Frame(self.comparison_display_frame)
        chart_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Создаем график сравнения
        self.comparison_fig, self.comparison_ax = plt.subplots(figsize=(6, 4))
        self.comparison_canvas = FigureCanvasTkAgg(self.comparison_fig, master=chart_frame)
        self.comparison_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Создаем фрейм для таблицы
        table_frame = ttk.Frame(self.comparison_display_frame)
        table_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Создаем таблицу сравнения
        columns = ("alternative", "eigenvector", "log", "geometric", "line", "ahp")
        self.comparison_tree = ttk.Treeview(table_frame, columns=columns, show="headings")

        # Настройка заголовков
        self.comparison_tree.heading("alternative", text="Альтернатива")
        self.comparison_tree.heading("eigenvector", text="Собственный\nвектор")
        self.comparison_tree.heading("log", text="Лог.\nнаим. кв.")
        self.comparison_tree.heading("geometric", text="Геометр.\nсреднее")
        self.comparison_tree.heading("line", text="Метод\n\"линия\"")
        self.comparison_tree.heading("ahp", text="Метод\nAHP")

        # Настройка колонок
        self.comparison_tree.column("alternative", width=120)
        self.comparison_tree.column("eigenvector", width=80, anchor=tk.E)
        self.comparison_tree.column("log", width=80, anchor=tk.E)
        self.comparison_tree.column("geometric", width=80, anchor=tk.E)
        self.comparison_tree.column("line", width=80, anchor=tk.E)
        self.comparison_tree.column("ahp", width=80, anchor=tk.E)

        # Добавляем скроллбар
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.comparison_tree.yview)
        self.comparison_tree.configure(yscrollcommand=scrollbar.set)
        self.comparison_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def update_comparison_tab(self):
        """Обновление вкладки сравнения методов."""
        if self.results is None:
            return

        # Обновляем график
        self.comparison_ax.clear()

        # Получаем имена альтернатив
        names = [var.get() for var in self.alternative_names]

        # Подготовка данных для графика
        weights_eig = self.results['weights_eigenvector']
        weights_log = self.results['weights_log_least_squares']
        weights_geo = self.results['weights_geometric_mean']
        weights_line = self.results['weights_line_method']
        weights_ahp, _, _, _ = calculate_weights_ahp(self.current_matrix)

        # Ширина столбцов
        bar_width = 0.15
        indices = np.arange(len(names))

        # Строим столбчатую диаграмму
        self.comparison_ax.bar(indices - 2 * bar_width, weights_eig, bar_width, label='Собственный вектор', alpha=0.8)
        self.comparison_ax.bar(indices - bar_width, weights_log, bar_width, label='Лог. наим. кв.', alpha=0.8)
        self.comparison_ax.bar(indices, weights_geo, bar_width, label='Геом. среднее', alpha=0.8)
        self.comparison_ax.bar(indices + bar_width, weights_line, bar_width, label='Метод "линия"', alpha=0.8)
        self.comparison_ax.bar(indices + 2 * bar_width, weights_ahp, bar_width, label='Метод AHP', alpha=0.8)

        self.comparison_ax.set_xlabel('Альтернативы')
        self.comparison_ax.set_ylabel('Веса')
        self.comparison_ax.set_title('Сравнение методов вычисления весов')
        self.comparison_ax.set_xticks(indices)
        self.comparison_ax.set_xticklabels(names, rotation=15)
        self.comparison_ax.legend()

        self.comparison_fig.tight_layout()
        self.comparison_canvas.draw()

        # Обновляем таблицу
        for item in self.comparison_tree.get_children():
            self.comparison_tree.delete(item)

        for i, name in enumerate(names):
            self.comparison_tree.insert("", tk.END, values=(
                name,
                f"{weights_eig[i]:.4f}",
                f"{weights_log[i]:.4f}",
                f"{weights_geo[i]:.4f}",
                f"{weights_line[i]:.4f}",
                f"{weights_ahp[i]:.4f}"
            ))

        self.log_message("Таблица сравнения методов обновлена.")

    def create_consistency_tab(self):
        """Создание вкладки для отображения показателей согласованности."""
        # Фрейм для основных показателей
        main_frame = ttk.Frame(self.consistency_frame)
        main_frame.pack(fill=tk.X, padx=5, pady=5)

        # Таблица основных показателей
        self.consistency_tree = ttk.Treeview(main_frame, columns=("metric", "value"), show="headings", height=8)
        self.consistency_tree.heading("metric", text="Показатель")
        self.consistency_tree.heading("value", text="Значение")
        self.consistency_tree.column("metric", width=250)
        self.consistency_tree.column("value", width=150, anchor=tk.E)
        self.consistency_tree.pack(fill=tk.X, padx=5, pady=5)

        # Фрейм для графика корреляции
        chart_frame = ttk.Frame(self.consistency_frame)
        chart_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем график корреляции
        self.corr_fig, self.corr_ax = plt.subplots(figsize=(6, 4))
        self.corr_canvas = FigureCanvasTkAgg(self.corr_fig, master=chart_frame)
        self.corr_canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

    def update_consistency_tab(self):
        """Обновление вкладки согласованности."""
        if self.results is None:
            return

        # Обновляем таблицу показателей
        for item in self.consistency_tree.get_children():
            self.consistency_tree.delete(item)

        metrics = [
            ("Собственное значение λ_max", f"{self.results['lambda_max']:.4f}"),
            ("Индекс согласованности CI", f"{self.results['CI']:.4f}"),
            ("Отношение согласованности CR", f"{self.results['CR']:.4f}"),
            ("Средняя корреляция", f"{self.results['mean_correlation']:.4f}"),
            ("Хи-квадрат статистика", f"{self.results['chi2_statistic']:.4f}"),
            ("p-значение для Хи-квадрат", f"{self.results['chi2_p_value']:.4f}"),
            ("Степени свободы", f"{self.results['chi2_df']}"),
            ("Размер матрицы", f"{self.current_matrix.shape[0]} x {self.current_matrix.shape[1]}")
        ]

        for metric, value in metrics:
            self.consistency_tree.insert("", tk.END, values=(metric, value))

        # Оценка согласованности
        cr_threshold = 0.1
        chi2_threshold = 0.05

        if self.results['CR'] < cr_threshold and self.results['chi2_p_value'] > chi2_threshold:
            self.log_message("Матрица парных сравнений считается согласованной (CR < 0.1 и p-value > 0.05).")
        else:
            self.log_message("Матрица парных сравнений имеет несогласованность (проверьте значения CR и p-value).")

        # Обновляем график корреляции
        self.corr_ax.clear()

        # Получаем матрицу корреляции
        corr_matrix = self.results['correlation_matrix']
        size = corr_matrix.shape[0]

        # Строим тепловую карту
        im = self.corr_ax.imshow(corr_matrix, cmap='coolwarm', vmin=-1, vmax=1)

        # Добавляем названия альтернатив
        names = [var.get() for var in self.alternative_names]
        self.corr_ax.set_xticks(range(size))
        self.corr_ax.set_yticks(range(size))
        self.corr_ax.set_xticklabels(names, rotation=45, ha='right')
        self.corr_ax.set_yticklabels(names)

        # Добавляем значения в ячейки
        for i in range(size):
            for j in range(size):
                self.corr_ax.text(j, i, f"{corr_matrix[i, j]:.2f}", ha="center", va="center",
                                   color="w" if abs(corr_matrix[i, j]) > 0.5 else "black")

        self.corr_ax.set_title('Матрица корреляций между альтернативами')
        self.corr_fig.colorbar(im, ax=self.corr_ax)
        self.corr_fig.tight_layout()
        self.corr_canvas.draw()

    def create_inconsistent_pairs_tab(self):
        """Создание вкладки для отображения несогласованных пар."""
        # Фрейм для фильтрации
        filter_frame = ttk.Frame(self.inconsistent_frame)
        filter_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(filter_frame, text="Фильтр по степени несогласованности >").pack(side=tk.LEFT)
        self.inconsistency_threshold = tk.DoubleVar(value=0.1)
        threshold_entry = ttk.Entry(filter_frame, textvariable=self.inconsistency_threshold, width=8)
        threshold_entry.pack(side=tk.LEFT, padx=(5, 10))

        ttk.Button(filter_frame, text="Фильтровать", command=self.update_inconsistent_pairs_tab).pack(side=tk.LEFT)
        ttk.Button(filter_frame, text="Сбросить", command=lambda: self.inconsistency_threshold.set(0.1)).pack(side=tk.LEFT, padx=(5, 0))

        # Таблица несогласованных пар
        self.inconsistent_tree = ttk.Treeview(self.inconsistent_frame, columns=("triplet", "inconsistency", "recommendation"), show="headings", height=10)
        self.inconsistent_tree.heading("triplet", text="Тройка альтернатив")
        self.inconsistent_tree.heading("inconsistency", text="Степень несогласованности")
        self.inconsistent_tree.heading("recommendation", text="Рекомендованное значение")
        self.inconsistent_tree.column("triplet", width=200)
        self.inconsistent_tree.column("inconsistency", width=150, anchor=tk.E)
        self.inconsistent_tree.column("recommendation", width=150, anchor=tk.E)
        self.inconsistent_tree.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Фрейм для рекомендаций
        rec_frame = ttk.LabelFrame(self.inconsistent_frame, text="Рекомендации")
        rec_frame.pack(fill=tk.X, padx=5, pady=5)

        self.recommendations_text = scrolledtext.ScrolledText(rec_frame, height=5, wrap=tk.WORD)
        self.recommendations_text.pack(fill=tk.X, padx=5, pady=5)
        self.recommendations_text.config(state=tk.DISABLED)

    def update_inconsistent_pairs_tab(self):
        """Обновление вкладки с несогласованными парами."""
        if self.results is None:
            return

        # Обновляем таблицу
        for item in self.inconsistent_tree.get_children():
            self.inconsistent_tree.delete(item)

        names = [var.get() for var in self.alternative_names]
        threshold = self.inconsistency_threshold.get()

        filtered_pairs = [(a, b, c, inc) for a, b, c, inc in self.results['inconsistent_pairs'] if inc > threshold]

        for i, (a, b, c, inc) in enumerate(filtered_pairs):
            triplet = f"{names[a]} → {names[b]} → {names[c]}"
            # Рассчитываем рекомендуемое значение
            recommended = self.current_matrix[a, b] * self.current_matrix[b, c]

            self.inconsistent_tree.insert("", tk.END, values=(triplet, f"{inc:.4f}", f"{recommended:.2f}"))

        # Обновляем рекомендации
        self.recommendations_text.config(state=tk.NORMAL)
        self.recommendations_text.delete(1.0, tk.END)

        if filtered_pairs:
            a, b, c, inc = filtered_pairs[0]

            names = [var.get() for var in self.alternative_names]

            rec_text = (
                f"Наиболее несогласованная тройка: {names[a]} → {names[b]} → {names[c]}\n\n"
                "Рекомендуется:\n"
                f"1. Проверить оценку сравнения {names[a]} и {names[b]} (значение: {self.current_matrix[a, b]:.2f})\n"
                f"2. Проверить оценку сравнения {names[b]} и {names[c]} (значение: {self.current_matrix[b, c]:.2f})\n"
                f"3. Проверить оценку сравнения {names[a]} и {names[c]} (значение: {self.current_matrix[a, c]:.2f})\n"
                f"4. Для согласованности значение {names[a]} и {names[c]} должно быть близко к {self.current_matrix[a, b] * self.current_matrix[b, c]:.2f}\n\n"
                "Возможные действия:\n"
                "- Пересмотреть экспертные оценки для указанных пар\n"
                "- Уточнить критерии сравнения\n"
                "- Провести дополнительное обсуждение с экспертами"
            )

            self.recommendations_text.insert(tk.END, rec_text)
        else:
            self.recommendations_text.insert(tk.END, "Несогласованные пары выше порога не обнаружены.")

        self.recommendations_text.config(state=tk.DISABLED)

    def create_3d_visualization_tab(self):
        """Создание вкладки для 3D визуализации."""
        # Фрейм для 3D графика
        self.vis_3d_frame = ttk.Frame(self.visualization_3d_frame)
        self.vis_3d_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Создаем 3D график
        self.fig_3d = plt.figure(figsize=(8, 6))
        self.ax_3d = self.fig_3d.add_subplot(111, projection='3d')
        self.canvas_3d = FigureCanvasTkAgg(self.fig_3d, master=self.vis_3d_frame)
        self.canvas_3d.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        # Фрейм для управления
        control_frame = ttk.Frame(self.visualization_3d_frame)
        control_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(control_frame, text="Обновить 3D график", command=self.update_3d_visualization).pack(side=tk.LEFT)
        ttk.Button(control_frame, text="Сменить вид", command=self.rotate_view).pack(side=tk.LEFT, padx=5)
        ttk.Button(control_frame, text="Экспорт 3D", command=self.export_3d_plot).pack(side=tk.LEFT, padx=5)

        # Инициализируем угол обзора
        self.elev = 30
        self.azim = 45

    def rotate_view(self):
        """Изменение угла обзора 3D графика."""
        self.elev = (self.elev + 30) % 360
        self.azim = (self.azim + 45) % 360
        self.update_3d_visualization()

    def export_3d_plot(self):
        """Экспорт 3D графика в файл."""
        if self.results is None:
            messagebox.showwarning("Экспорт", "Сначала выполните анализ матрицы.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG files", "*.png"), ("PDF files", "*.pdf"), ("SVG files", "*.svg")],
            title="Экспорт 3D графика"
        )

        if file_path:
            try:
                self.fig_3d.savefig(file_path, dpi=300, bbox_inches='tight')

                self.log_message(f"3D график экспортирован в файл: {os.path.basename(file_path)}")
                messagebox.showinfo("Экспорт", f"3D график успешно сохранен в файл:\n{file_path}")

            except Exception as e:
                messagebox.showerror("Ошибка экспорта", f"Не удалось экспортировать 3D график: {str(e)}")
                self.log_message(f"Ошибка экспорта 3D графика: {str(e)}")

    def update_3d_visualization(self):
        """Обновление 3D визуализации."""
        if self.results is None:
            return

        # Очищаем предыдущий график
        self.ax_3d.clear()

        # Получаем имена альтернатив
        names = [var.get() for var in self.alternative_names]
        n = len(names)

        # Создаем данные для 3D графика
        x_pos = np.arange(n)
        y_pos = np.arange(n)
        x_pos, y_pos = np.meshgrid(x_pos, y_pos)

        # Используем матрицу парных сравнений как z-значения
        z_values = self.current_matrix

        # Строим поверхность
        surf = self.ax_3d.plot_surface(x_pos, y_pos, z_values, cmap='viridis', alpha=0.8)

        # Добавляем бары для весов
        weights = self.results['weights_eigenvector']

        for i, w in enumerate(weights):
            self.ax_3d.bar3d(i, n + 1, 0, 0.5, 0.5, w, color=self.color_map(i % 12), alpha=0.8)

        # Настройка осей
        self.ax_3d.set_xlabel('Альтернативы (столбцы)')
        self.ax_3d.set_ylabel('Альтернативы (строки)')
        self.ax_3d.set_zlabel('Значение')

        # Установка меток
        self.ax_3d.set_xticks(range(n))
        self.ax_3d.set_xticklabels(names, rotation=45)
        self.ax_3d.set_yticks(range(n))
        self.ax_3d.set_yticklabels(names, rotation=-45)

        # Добавляем легенду для весов
        self.ax_3d.text2D(0.05, 0.95, "Веса альтернатив:", transform=self.ax_3d.transAxes, fontsize=10)
        for i, (name, w) in enumerate(zip(names, weights)):
            self.ax_3d.text2D(0.05, 0.90 - i * 0.05, f"{name}: {w:.3f}", transform=self.ax_3d.transAxes, fontsize=9)

        # Установка угла обзора
        self.ax_3d.view_init(elev=self.elev, azim=self.azim)

        # Добавляем цветовую шкалу
        self.fig_3d.colorbar(surf, ax=self.ax_3d, shrink=0.5, aspect=5)
        self.ax_3d.set_title('3D Визуализация матрицы парных сравнений и весов альтернатив')
        self.fig_3d.tight_layout()
        self.canvas_3d.draw()

    def create_log_tab(self):
        """Создание вкладки для лога выполнения."""
        self.log_text = scrolledtext.ScrolledText(self.log_frame, wrap=tk.WORD, state=tk.DISABLED)
        self.log_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Кнопки управления логом
        btn_frame = ttk.Frame(self.log_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        ttk.Button(btn_frame, text="Очистить лог", command=self.clear_log).pack(side=tk.LEFT)
        ttk.Button(btn_frame, text="Скопировать лог", command=self.copy_log).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(btn_frame, text="Сохранить лог в файл", command=self.save_log).pack(side=tk.LEFT, padx=(5, 0))
        ttk.Button(btn_frame, text="Экспорт полного отчета", command=self.export_full_report).pack(side=tk.RIGHT, padx=(5, 0))

        # Добавляем кнопки отката
        undo_redo_frame = ttk.Frame(self.log_frame)
        undo_redo_frame.pack(fill=tk.X, padx=5, pady=(0, 5))

        ttk.Button(undo_redo_frame, text="Undo (Ctrl+Z)", command=self.undo_operation).pack(side=tk.LEFT)
        ttk.Button(undo_redo_frame, text="Redo (Ctrl+Y)", command=self.redo_operation).pack(side=tk.LEFT, padx=(5, 0))

        # Привязываем сочетания клавиш
        self.root.bind('<Control-z>', lambda e: self.undo_operation())
        self.root.bind('<Control-y>', lambda e: self.redo_operation())

    def log_message(self, message):
        """Добавление сообщения в лог."""
        try:
            timestamp = datetime.now().strftime('%H:%M:%S')
            self.log_text.config(state=tk.NORMAL)
            self.log_text.insert(tk.END, f"[{timestamp}] {message}\n")
            self.log_text.see(tk.END)
            self.log_text.config(state=tk.DISABLED)
        except Exception as e:
            print(f"[LOG ERROR] {message} | Exception: {e}")

    def clear_log(self):
        """Очистка лога."""
        self.log_text.config(state=tk.NORMAL)
        self.log_text.delete(1.0, tk.END)
        self.log_text.config(state=tk.DISABLED)
        self.log_message("Лог очищен.")

    def copy_log(self):
        """Копирование лога в буфер обмена."""
        log_content = self.log_text.get(1.0, tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(log_content)
        self.log_message("Лог скопирован в буфер обмена.")

    def save_log(self):
        """Сохранение лога в файл."""
        log_content = self.log_text.get(1.0, tk.END)
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Текстовые файлы", "*.txt"), ("Все файлы", "*.*")],
            title="Сохранить лог"
        )

        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(log_content)
                self.log_message(f"Лог сохранен в файл: {os.path.basename(file_path)}")
            except Exception as e:
                messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить лог: {str(e)}")
                self.log_message(f"Ошибка при сохранении лога: {str(e)}")

    def add_to_operation_history(self, operation):
        """Добавление операции в историю."""
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.operation_history.append({'timestamp': timestamp, 'operation': operation})

        # Ограничиваем размер истории
        if len(self.operation_history) > self.max_history_size:
            self.operation_history.pop(0)

    def undo_operation(self):
        """Откат последней операции"""
        if not self.operation_history:
            messagebox.showinfo("Undo", "Нет операций для отката.")
            return

        last_operation = self.operation_history.pop()
        self.log_message(f"Откат операции: {last_operation['operation']} (время: {last_operation['timestamp']})")

    def redo_operation(self):
        """Повтор отмененной операции."""
        messagebox.showinfo("Redo", "Функция Redo пока не реализована полностью.")

    def export_full_report(self):
        """Экспорт полного отчета по лабораторной работе."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Экспорт отчета", "Сначала выполните анализ матрицы.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[
                ("Word документы", "*.docx"),
                ("PDF файлы", "*.pdf"),
                ("Текстовые файлы", "*.txt")
            ],
            title="Экспорт полного отчета"
        )

        if not file_path:
            return

        try:
            # Создаем отчет
            report = self.generate_report()

            # Сохраняем в зависимости от расширения
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.docx':
                # Экспорт в Word документ
                self.export_to_docx(report, file_path)
            elif ext == '.pdf':
                # Экспорт в PDF
                self.export_to_pdf(report, file_path)
            else:  # .txt
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(report)

            self.log_message(f"Полный отчет экспортирован в файл: {os.path.basename(file_path)}")
            messagebox.showinfo("Экспорт отчета", f"Полный отчет успешно сохранен в файл:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка экспорта", f"Не удалось экспортировать отчет: {str(e)}")
            self.log_message(f"Ошибка при экспорте отчета: {str(e)}")

    def generate_report(self):
        """Генерация полного отчета по лабораторной работе."""
        report = []
        report.append("=" * 60)
        report.append("ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №7")
        report.append("Поддержка принятия решений и повышение согласованности экспертных оценок")
        report.append("Вариант №1: Разработка информационных систем")
        report.append("=" * 60)
        report.append(f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        report.append("\n")

        # Исходная матрица
        report.append("1. ИСХОДНАЯ МАТРИЦА ПАРНЫХ СРАВНЕНИЙ")
        report.append("-" * 40)

        names = [var.get() for var in self.alternative_names]
        report.append(" " + " ".join([f"{name:>10}" for name in names]))

        for i, name in enumerate(names):
            row = [f"{name:>10}"]
            for j in range(len(names)):
                if i == j:
                    row.append(" 1.00")
                else:
                    row.append(f"{self.current_matrix[i][j]:>10.2f}")
            report.append("".join(row))

        report.append("\n")

        # Веса альтернатив
        report.append("2. ВЕСА АЛЬТЕРНАТИВ")
        report.append("-" * 40)
        report.append(
            f"{'Альтернатива':<15} {'Собственный вектор':<15} {'Лог. наим. кв.':<15} {'Геом. среднее':<15} {'Метод \"линия\"':<15} {'Метод AHP':<15}")

        weights_eig = self.results['weights_eigenvector']
        weights_log = self.results['weights_log_least_squares']
        weights_geo = self.results['weights_geometric_mean']
        weights_line = self.results['weights_line_method']
        weights_ahp, _, _, _ = calculate_weights_ahp(self.current_matrix)

        for i, name in enumerate(names):
            report.append(
                f"{name:<15} {weights_eig[i]:<15.6f} {weights_log[i]:<15.6f} {weights_geo[i]:<15.6f} {weights_line[i]:<15.6f} {weights_ahp[i]:<15.6f}"
            )

        report.append("\n")

        # Показатели согласованности
        report.append("3. ПОКАЗАТЕЛИ СОГЛАСОВАННОСТИ")
        report.append("-" * 40)

        report.append(f"Собственное значение λ_max: {self.results['lambda_max']:.4f}")
        report.append(f"Индекс согласованности CI: {self.results['CI']:.4f}")
        report.append(f"Отношение согласованности CR: {self.results['CR']:.4f}")
        report.append(f"Средняя корреляция: {self.results['mean_correlation']:.4f}")
        report.append(f"Хи-квадрат статистика: {self.results['chi2_statistic']:.4f}")
        report.append(f"р-значение для Хи-квадрат: {self.results['chi2_p_value']:.4f}")
        report.append(f"Степени свободы: {self.results['chi2_df']}")

        # Оценка согласованности
        cr_threshold = 0.1
        chi2_threshold = 0.05
        if self.results['CR'] < cr_threshold and self.results['chi2_p_value'] > chi2_threshold:
            report.append("Оценка: Матрица считается согласованной.")
        else:
            report.append("Оценка: Матрица имеет несогласованность.")

        report.append("\n")

        # Несогласованные пары
        report.append("4. НЕСОГЛАСОВАННЫЕ ПАРЫ")
        report.append("-" * 40)

        if self.results['inconsistent_pairs']:
            report.append(f"{'Тройка альтернатив':<30} {'Несогласованность':<15} {'Рекомендованное значение'}")
            for a, b, c, inc in self.results['inconsistent_pairs'][:5]:
                triplet = f"{names[a]}→{names[b]}→{names[c]}"
                recommended = self.current_matrix[a, b] * self.current_matrix[b, c]
                report.append(f"{triplet:<30} {inc:<15.4f} {recommended:<20.2f}")
        else:
            report.append("Несогласованные пары не обнаружены.")

        report.append("\n")

        # Анализ чувствительности
        report.append("5. АНАЛИЗ ЧУВСТВИТЕЛЬНОСТИ")
        report.append("-" * 40)

        if hasattr(self, 'sensitivity_results') and self.sensitivity_results:
            i, j = self.sensitivity_results['most_sensitive_pair']
            report.append(f"Наиболее чувствительная пара: {names[i]} vs {names[j]}")
            report.append(f"Максимальная чувствительность: {self.sensitivity_results['max_sensitivity']:.6f}")
            report.append(f"Фактор возмущения: {self.sensitivity_results['perturbation_factor'] * 100:.1f}%")
        else:
            report.append("Анализ чувствительности не выполнен.")

        report.append("\n")

        # Выводы
        report.append("6. ВЫВОДЫ")
        report.append("-" * 40)

        best_idx_eig = np.argmax(self.results['weights_eigenvector'])
        best_idx_line = np.argmax(self.results['weights_line_method'])
        best_idx_ahp = np.argmax(weights_ahp)

        report.append(f"По методу собственного вектора лучшая альтернатива: {names[best_idx_eig]}")
        report.append(f"По методу \"линия\" лучшая альтернатива: {names[best_idx_line]}")
        report.append(f"По методу AHP лучшая альтернатива: {names[best_idx_ahp]}")

        if best_idx_eig == best_idx_line == best_idx_ahp:
            report.append("Разные методы дают одинаковый результат.")
        else:
            report.append("Разные методы дают различный результат.")

        cr_threshold = 0.1
        if self.results['CR'] < cr_threshold:
            report.append("Матрица является согласованной (CR < 0.1).")
        else:
            report.append("Матрица не является полностью согласованной (CR >= 0.1).")

        report.append("\n")
        report.append("Автор: Колесов Станислав")
        report.append("Дата: " + datetime.now().strftime("%Y-%m-%d"))

        return "\n".join(report)

    def export_to_docx(self, report, file_path):
        """Экспорт отчета в Word документ."""
        try:
            doc = Document()

            # Добавляем заголовок
            title = doc.add_heading('ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №7', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавляем подзаголовок
            subtitle = doc.add_paragraph('Поддержка принятия решений и повышение согласованности экспертных оценок')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.bold = True

            # Добавляем информацию о варианте
            variant_info = doc.add_paragraph('Вариант №1: Разработка информационных систем')
            variant_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            variant_info.runs[0].font.bold = True

            # Добавляем дату
            date_info = doc.add_paragraph(f'Дата: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Добавляем разделы отчета
            sections = report.split('\n\n')
            for section in sections:
                if section.strip():
                    # Если секция начинается с цифры и точки, это заголовок
                    if section.strip()[0].isdigit() and '.' in section.strip():
                        heading = doc.add_heading(section.strip(), level=1)
                    else:
                        # Иначе добавляем как обычный абзац
                        paragraph = doc.add_paragraph(section.strip())

            # Сохраняем документ
            doc.save(file_path)

        except Exception as e:
            raise e

    def export_to_pdf(self, report, file_path):
        """Экспорт отчета в PDF."""
        try:
            # Создаем PDF документ
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Добавляем заголовок
            c.setFont("Helvetica-Bold", 16)
            c.drawString(1 * inch, height - 1 * inch, "ОТЧЕТ ПО ЛАБОРАТОРНОЙ РАБОТЕ №7")

            # Добавляем подзаголовок
            c.setFont("Helvetica-Bold", 12)
            c.drawString(1 * inch, height - 1.5 * inch, "Поддержка принятия решений и повышение согласованности экспертных оценок")

            # Добавляем информацию о варианте
            c.setFont("Helvetica-Bold", 12)
            c.drawString(1 * inch, height - 1.8 * inch, "Вариант №1: Разработка информационных систем")

            # Добавляем дату
            c.setFont("Helvetica", 10)
            c.drawString(1 * inch, height - 2.1 * inch, f"Дата: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # Добавляем содержание отчета
            c.setFont("Helvetica", 10)

            y_position = height - 2.5 * inch

            # Разделяем отчет на строки
            lines = report.split('\n')

            for line in lines:
                if y_position < 1 * inch:
                    c.showPage()
                    y_position = height - 1 * inch

                c.drawString(1 * inch, y_position, line)
                y_position -= 12

            # Сохраняем PDF
            c.save()

        except Exception as e:
            raise e

    def export_results(self):
        """Экспорт результатов анализа в файл."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Экспорт результатов", "Сначала выполните анализ матрицы.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[
                ("Excel файлы", "*.xlsx"),
                ("CSV файлы", "*.csv"),
                ("Текстовые файлы", "*.txt"),
                ("LaTeX файлы", "*.tex")
            ],
            title="Экспортировать результаты"
        )

        if not file_path:
            return

        try:
            # Создаем DataFrame с результатами
            size = self.current_matrix.shape[0]
            names = [var.get() for var in self.alternative_names]

            # Таблица весов
            weights_data = {
                "Альтернатива": names,
                "Метод собственного вектора": self.results['weights_eigenvector'],
                "Метод лог. наим. кв.": self.results['weights_log_least_squares'],
                "Метод геометрического среднего": self.results['weights_geometric_mean'],
                "Метод \"линия\"": self.results['weights_line_method'],
                "Метод AHP": calculate_weights_ahp(self.current_matrix)[0]
            }
            weights_df = pd.DataFrame(weights_data)

            # Показатели согласованности
            consistency_data = {
                "Показатель": [
                    "Собственное значение λ_max",
                    "Индекс согласованности CI",
                    "Отношение согласованности CR",
                    "Средняя корреляция",
                    "Хи-квадрат статистика",
                    "p-значение для Хи-квадрат",
                    "Степени свободы",
                    "Размер матрицы"
                ],
                "Значение": [
                    f"{self.results['lambda_max']:.4f}",
                    f"{self.results['CI']:.4f}",
                    f"{self.results['CR']:.4f}",
                    f"{self.results['mean_correlation']:.4f}",
                    f"{self.results['chi2_statistic']:.4f}",
                    f"{self.results['chi2_p_value']:.4f}",
                    f"{self.results['chi2_df']}",
                    f"{size}x{size}"
                ]
            }
            consistency_df = pd.DataFrame(consistency_data)

            # Несогласованные пары
            inconsistent_data = []
            for i, (a, b, c, inc) in enumerate(self.results['inconsistent_pairs']):
                inconsistent_data.append({
                    "Тройка альтернатив": f"{names[a]} → {names[b]} → {names[c]}",
                    "Степень несогласованности": f"{inc:.4f}",
                    "Рекомендованное значение": f"{self.current_matrix[a, b] * self.current_matrix[b, c]:.2f}"
                })
            inconsistent_df = pd.DataFrame(inconsistent_data)

            # Исходная матрица
            matrix_df = pd.DataFrame(self.current_matrix, index=names, columns=names)

            # Сохраняем в зависимости от расширения
            ext = os.path.splitext(file_path)[1].lower()

            if ext == '.xlsx':
                with pd.ExcelWriter(file_path) as writer:
                    weights_df.to_excel(writer, sheet_name='Веса', index=False)
                    consistency_df.to_excel(writer, sheet_name='Согласованность', index=False)
                    inconsistent_df.to_excel(writer, sheet_name='Несогласованные пары', index=False)
                    matrix_df.to_excel(writer, sheet_name='Матрица')
            elif ext == '.csv':
                # Для CSV сохраняем только основные результаты
                output = io.StringIO()
                weights_df.to_csv(output, index=False)
                output.write("\n\nПоказатели согласованности:\n")
                consistency_df.to_csv(output, index=False)
                output.write("\n\nНесогласованные пары:\n")
                inconsistent_df.to_csv(output, index=False)

                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(output.getvalue())
            elif ext == '.tex':
                # Экспорт в LaTeX
                latex_weights = export_to_latex_table(weights_df, "Веса альтернатив", "tab:weights")
                latex_consistency = export_to_latex_table(consistency_df, "Показатели согласованности", "tab:consistency")

                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("\\documentclass{article}\n")
                    f.write("\\usepackage[utf8]{inputenc}\n")
                    f.write("\\usepackage[russian]{babel}\n")
                    f.write("\\usepackage{amsmath}\n")
                    f.write("\\usepackage{graphicx}\n")
                    f.write("\\usepackage{array}\n")
                    f.write("\\begin{document}\n")
                    f.write("\\title{Отчет по лабораторной работе №7}\n")
                    f.write("\\author{Колосов Станислав}\n")
                    f.write("\\date{" + datetime.now().strftime("%Y-%m-%d") + "}\n")
                    f.write("\\maketitle\n")
                    f.write("\\section{Веса альтернатив}\n")
                    f.write(latex_weights)
                    f.write("\\section{Показатели согласованности}\n")
                    f.write(latex_consistency)
                    f.write("\\end{document}\n")

            else:  # .txt
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("===== ВЕСА АЛЬТЕРНАТИВ =====\n")
                    for _, row in weights_df.iterrows():
                        f.write(f"{row['Альтернатива']}: "
                                f"Собственный вектор={row['Метод собственного вектора']:.6f}, "
                                f"Лог. наим. кв.={row['Метод лог. наим. кв.']:.6f}, "
                                f"Геом. среднее={row['Метод геометрического среднего']:.6f}, "
                                f"Метод \"линия\"={row['Метод \"линия\"']:.6f}, "
                                f"Метод AHP={row['Метод AHP']:.6f}\n")

                    f.write("\n===== ПОКАЗАТЕЛИ СОГЛАСОВАННОСТИ =====\n")
                    for _, row in consistency_df.iterrows():
                        f.write(f"{row['Показатель']}: {row['Значение']}\n")

                    f.write("\n===== НЕСОГЛАСОВАННЫЕ ПАРЫ =====\n")
                    for _, row in inconsistent_df.iterrows():
                        f.write(
                            f"{row['Тройка альтернатив']}: {row['Степень несогласованности']} (рекомендовано: {row['Рекомендованное значение']}) \n")
                    f.write("\n===== ИСХОДНАЯ МАТРИЦА ======\n")
                    for i, name in enumerate(names):
                        f.write(f"{name}: " + ", ".join([f"{self.current_matrix[i, j]:.2f}" for j in range(size)]) + "\n")

            self.log_message(f"Результаты экспортированы в файл: {os.path.basename(file_path)}")
            messagebox.showinfo("Экспорт результатов", f"Результаты успешно сохранены в файл:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка экспорта", f"Не удалось экспортировать результаты: {str(e)}")
            self.log_message(f"Ошибка при экспорте результатов: {str(e)}")

    def save_session(self):
        """Сохранение сессии анализа."""
        if self.current_matrix is None or self.results is None:
            messagebox.showwarning("Сохранение сессии", "Сначала выполните анализ матрицы.")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON файлы", "*.json"), ("Все файлы", "*.*")],
            title="Сохранить сессию"
        )

        if not file_path:
            return

        try:
            # Подготовка данных сессии
            session_data = {
                "timestamp": datetime.now().isoformat(),
                "matrix_size": self.matrix_size.get(),
                "alternative_names": [var.get() for var in self.alternative_names],
                "matrix": self.current_matrix.tolist(),
                "results": {
                    "weights_eigenvector": self.results['weights_eigenvector'].tolist(),
                    "weights_log_least_squares": self.results['weights_log_least_squares'].tolist(),
                    "weights_geometric_mean": self.results['weights_geometric_mean'].tolist(),
                    "weights_line_method": self.results['weights_line_method'].tolist(),
                    "lambda_max": float(self.results['lambda_max']),
                    "CI": float(self.results['CI']),
                    "CR": float(self.results['CR']),
                    "mean_correlation": float(self.results['mean_correlation']),
                    "chi2_statistic": float(self.results['chi2_statistic']),
                    "chi2_p_value": float(self.results['chi2_p_value']),
                    "chi2_df": int(self.results['chi2_df']),
                    "inconsistent_pairs": [[int(a), int(b), int(c), float(inc)] for a, b, c, inc in self.results['inconsistent_pairs']],
                    "consistent_matrix": self.results['consistent_matrix'].tolist()
                },
                "sensitivity_results": None
            }

            if hasattr(self, 'sensitivity_results') and self.sensitivity_results:
                session_data["sensitivity_results"] = {
                    "max_sensitivity": float(self.sensitivity_results['max_sensitivity']),
                    "most_sensitive_pair": [int(i) for i in self.sensitivity_results['most_sensitive_pair']],
                    "perturbation_factor": float(self.sensitivity_results['perturbation_factor'])
                }

            # Сохраняем сессию
            success = save_session(session_data, file_path)

            if success:
                self.session_file = file_path
                self.log_message(f"Сессия сохранена в файл: {os.path.basename(file_path)}")
                messagebox.showinfo("Сохранение сессии", f"Сессия успешно сохранена в файл:\n{file_path}")
            else:
                raise Exception("Не удалось сохранить сессию")

        except Exception as e:
            messagebox.showerror("Ошибка сохранения", f"Не удалось сохранить сессию: {str(e)}")
            self.log_message(f"Ошибка сохранения сессии: {str(e)}")

    def load_session(self):
        """Загрузка сессии анализа."""
        file_path = filedialog.askopenfilename(
            defaultextension=".json",
            filetypes=[("JSON файлы", "*.json"), ("Все файлы", "*.*")],
            title="Загрузить сессию"
        )

        if not file_path:
            return

        try:
            # Загружаем сессию
            session_data = load_session(file_path)
            if session_data is None:
                raise Exception("Не удалось загрузить сессию")

            # Восстанавливаем состояние приложения
            self.matrix_size.set(session_data["matrix_size"])
            self.create_matrix_entries(session_data["matrix_size"])

            # Устанавливаем имена альтернатив
            for i, name in enumerate(session_data["alternative_names"]):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Восстанавливаем матрицу
            matrix_data = np.array(session_data["matrix"])
            self.current_matrix = matrix_data

            # Восстанавливаем результаты
            results_data = session_data["results"]
            self.results = {
                'weights_eigenvector': np.array(results_data["weights_eigenvector"]),
                'weights_log_least_squares': np.array(results_data["weights_log_least_squares"]),
                'weights_geometric_mean': np.array(results_data["weights_geometric_mean"]),
                'weights_line_method': np.array(results_data["weights_line_method"]),
                'lambda_max': results_data["lambda_max"],
                'CI': results_data["CI"],
                'CR': results_data["CR"],
                'mean_correlation': results_data["mean_correlation"],
                'chi2_statistic': results_data["chi2_statistic"],
                'chi2_p_value': results_data["chi2_p_value"],
                'chi2_df': results_data["chi2_df"],
                'inconsistent_pairs': [(a, b, c, inc) for a, b, c, inc in results_data["inconsistent_pairs"]],
                'consistent_matrix': np.array(results_data["consistent_matrix"])
            }

            # Восстанавливаем анализ чувствительности, если есть
            if "sensitivity_results" in session_data and session_data["sensitivity_results"]:
                self.sensitivity_results = {
                    'max_sensitivity': session_data["sensitivity_results"]["max_sensitivity"],
                    'most_sensitive_pair': tuple(session_data["sensitivity_results"]["most_sensitive_pair"]),
                    'perturbation_factor': session_data["sensitivity_results"]["perturbation_factor"]
                }

            # Заполняем матрицу в интерфейсе
            for i, row in enumerate(matrix_data):
                for j, val in enumerate(row):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{val:.2f}")

            # Обновляем все вкладки
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.session_file = file_path
            self.log_message(f"Сессия загружена из файла: {os.path.basename(file_path)}")
            messagebox.showinfo("Загрузка сессии", f"Сессия успешно загружена из файла:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка загрузки", f"Не удалось загрузить сессию: {str(e)}")
            self.log_message(f"Ошибка загрузки сессии: {str(e)}")

    def import_from_csv(self):
        """Импорт матрицы из CSV файла."""
        file_path = filedialog.askopenfilename(
            defaultextension=".csv",
            filetypes=[("CSV файлы", "*.csv"), ("Все файлы", "*.*")],
            title="Импорт из CSV"
        )

        if not file_path:
            return

        try:
            # Загружаем данные из CSV
            df = pd.read_csv(file_path, header=None)
            # Проверяем размеры
            if df.shape[0] != df.shape[1]:
                raise ValueError("Матрица должна быть квадратной")
            # Устанавливаем размер матрицы
            size = df.shape[0]
            self.matrix_size.set(size)
            self.create_matrix_entries(size)
            # Устанавливаем имена альтернатив (если они есть в CSV)
            if len(df.columns) > size:
                # Первый столбец содержит имена строк
                row_names = df.iloc[:, 0].astype(str).tolist()
                col_names = df.iloc[0, :].astype(str).tolist()
                data_start_row = 1
                data_start_col = 1
            else:
                # Только данные, без имен
                row_names = [f"Альтернатива {i + 1}" for i in range(size)]
                col_names = [f"Альтернатива {j + 1}" for j in range(size)]
                data_start_row = 0
                data_start_col = 0

            # Устанавливаем имена альтернатив
            for i, name in enumerate(row_names):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Заполняем матрицу данными
            matrix_data = df.iloc[data_start_row: data_start_row + size, data_start_col: data_start_col + size].values.astype(float)

            # Проверяем, что матрица положительно определенная
            for i in range(size):
                for j in range(size):
                    if i == j:
                        if matrix_data[i, j] != 1:
                            raise ValueError(f"Диагональный элемент [{i}, {j}] должен быть 1, найдено: {matrix_data[i, j]}")
                    else:
                        if matrix_data[i, j] <= 0:
                            raise ValueError(f"Элемент [{i}, {j}] должен быть положительным, найдено: {matrix_data[i, j]}")
                        if abs(matrix_data[i, j] * matrix_data[j, i] - 1) > 1e-10:
                            raise ValueError(f"Элементы [{i}, {j}] и [{j}, {i}] не являются обратными друг другу")

            # Устанавливаем матрицу
            self.current_matrix = matrix_data

            # Заполняем интерфейс
            for i, row in enumerate(matrix_data):
                for j, val in enumerate(row):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{val:.2f}")

            # Выполняем анализ
            self.results = analyze_consistency(matrix_data)
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.log_message(f"Матрица успешно импортирована из CSV файла: {os.path.basename(file_path)}")
            messagebox.showinfo("Импорт из CSV", f"Матрица успешно импортирована из файла:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка импорта", f"Не удалось импортировать матрицу из CSV: {str(e)}")
            self.log_message(f"Ошибка импорта из CSV: {str(e)}")

    def import_from_excel(self):
        """Импорт матрицы из Excel файла."""
        file_path = filedialog.askopenfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")],
            title="Импорт из Excel"
        )

        if not file_path:
            return

        try:
            # Загружаем данные из Excel
            excel_file = pd.ExcelFile(file_path)

            # Предлагаем выбрать лист
            if len(excel_file.sheet_names) > 1:
                sheet_name = simpledialog.askstring("Выбор листа", f"Выберите лист:\n{', '.join(excel_file.sheet_names)}", parent=self.root)
                if sheet_name not in excel_file.sheet_names:
                    sheet_name = excel_file.sheet_names[0]
            else:
                sheet_name = excel_file.sheet_names[0]

            # Загружаем выбранный лист
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

            # Проверяем размеры
            if df.shape[0] != df.shape[1]:
                raise ValueError("Матрица должна быть квадратной")

            # Устанавливаем размер матрицы
            size = df.shape[0]
            self.matrix_size.set(size)
            self.create_matrix_entries(size)

            # Устанавливаем имена альтернатив (если они есть в Excel)
            if len(df.columns) > size:
                # Первый столбец содержит имена строк
                row_names = df.iloc[:, 0].astype(str).tolist()
                col_names = df.iloc[0, :].astype(str).tolist()
                data_start_row = 1
                data_start_col = 1
            else:
                # Только данные, без имен
                row_names = [f"Альтернатива {i + 1}" for i in range(size)]
                col_names = [f"Альтернатива {j + 1}" for j in range(size)]
                data_start_row = 0
                data_start_col = 0

            # Устанавливаем имена альтернатив
            for i, name in enumerate(row_names):
                if i < len(self.alternative_names):
                    self.alternative_names[i].set(name)

            # Заполняем матрицу данными
            matrix_data = df.iloc[data_start_row: data_start_row + size, data_start_col: data_start_col + size].values.astype(float)

            # Проверяем, что матрица положительно определенная
            for i in range(size):
                for j in range(size):
                    if i == j:
                        if matrix_data[i, j] != 1:
                            raise ValueError(f"Диагональный элемент [{i}, {j}] должен быть 1, найдено: {matrix_data[i, j]}")
                    else:
                        if matrix_data[i, j] <= 0:
                            raise ValueError(f"Элемент [{i}, {j}] должен быть положительным, найдено: {matrix_data[i, j]}")
                        if abs(matrix_data[i, j] * matrix_data[j, i] - 1) > 1e-10:
                            raise ValueError(f"Элементы [{i}, {j}] и [{j}, {i}] не являются обратными друг другу")

            # Устанавливаем матрицу
            self.current_matrix = matrix_data

            # Заполняем интерфейс
            for i in range(size):
                for j in range(size):
                    if i != j:
                        self.matrix_entries[i][j].set(f"{matrix_data[i, j]:.2f}")

            # Выполняем анализ
            self.results = analyze_consistency(matrix_data)
            self.update_weights_tab()
            self.update_comparison_tab()
            self.update_consistency_tab()
            self.update_inconsistent_pairs_tab()
            self.log_message(f"Матрица успешно импортирована из Excel файла: {os.path.basename(file_path)}")
            messagebox.showinfo("Импорт из Excel", f"Матрица успешно импортирована из файла:\n{file_path}")

        except Exception as e:
            messagebox.showerror("Ошибка импорта", f"Не удалось импортировать матрицу из Excel: {str(e)}")
            self.log_message(f"Ошибка импорта из Excel: {str(e)}")

    def show_about(self):
        """Показывает информацию о программе."""
        about_text = (
            "Программа для поддержки принятия решений и повышения согласованности экспертных оценок\n"
            "Лабораторная работа №7 по дисциплине \"Основы ИИ\"\n"
            "Вариант 1: Разработка информационных систем\n\n"

            "Основные функции:\n"
            "- Анализ матрицы парных сравнений\n"
            "- Вычисление весов альтернатив четырьмя методами\n"
            "- Оценка согласованности экспертных оценок\n"
            "- Выявление наиболее несогласованных парных сравнений\n"
            "- Анализ чувствительности и корректировка несогласованности\n"
            "- 3D визуализация результатов\n"
            "- Экспорт полного отчета\n\n"

            "Разработано для учебных целей\n"
            "Автор: Колосов Станислав"
        )

        messagebox.showinfo("О программе", about_text)

    def toggle_theme(self):
        """Переключение цветовой темы интерфейса."""
        if self.current_theme == "light":
            self.current_theme = "dark"
            self.root.config(bg=self.themes["dark"]["bg"])
            self.style.configure("TFrame", background=self.themes["dark"]["bg"])
            self.style.configure("TLabel", background=self.themes["dark"]["bg"], foreground=self.themes["dark"]["fg"])
            self.style.configure("Header.TLabel", background=self.themes["dark"]["bg"], foreground=self.themes["dark"]["fg"])
        else:
            self.current_theme = "light"
            self.root.config(bg=self.themes["light"]["bg"])
            self.style.configure("TFrame", background=self.themes["light"]["bg"])
            self.style.configure("TLabel", background=self.themes["light"]["bg"], foreground=self.themes["light"]["fg"])
            self.style.configure("Header.TLabel", background=self.themes["light"]["bg"], foreground=self.themes["light"]["fg"])

        self.log_message(f"Тема интерфейса изменена на {self.current_theme}.")

if __name__ == "__main__":
    root = tk.Tk()
    app = DecisionSupportApp(root)
    root.mainloop()
